import os
import pytz 
import base64
import re
import json
import sqlite3
import threading
import asyncio
import httpx
import functools
import logging
import schedule
import time
import xml.etree.ElementTree as ET
import random
from datetime import datetime, timedelta
from docx import Document
from gtts import gTTS
from difflib import SequenceMatcher
from flask import Flask, request, Response, send_file
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
    CallbackQueryHandler,
)

# ====== Logging Configuration ======
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("bot_errors.log"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)

# ====== Configuration ======
TELEGRAM_BOT_TOKEN = os.environ.get("BOT_TOKEN") or "YOUR_BOT_TOKEN_HERE"
OPENROUTER_API_KEY = os.environ.get("OR_KEY") or "YOUR_OPENROUTER_KEY_HERE"
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

ADMIN_IDS = [1119128553]
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "mysecret123")

ADMIN_USERS = {
    "admin": "mysecret123",
    "thawkhyan": "yourpass123",
}

EXCHANGE_RATE = 4545

PAYMENT_INFO = (
    "💳 **ငွေလွှဲရန်**\n"
    "KBZPay: 09426419462\n"
    "WavePay: 09426419462"
)

PLAN_LIMITS = {
    "free": {
        "limit": 80,
        "price": 0,
        "prices": {"1m": 0, "3m": 0, "12m": 0}
    },
    "basic": {
        "limit": 1000,
        "price": 7000,        # 10000 ကနေ 7000 သို့ လျှော့
        "prices": {"1m": 7000, "3m": 17500, "12m": 56000}
    },
    "premium": {
        "limit": 3000,
        "price": 20000,       # 30000 ကနေ 20000 သို့ လျှော့
        "prices": {"1m": 20000, "3m": 50000, "12m": 160000}
    },
    "premium_plus": {
        "limit": 10000,
        "price": 35000,       # 50000 ကနေ 35000 သို့ လျှော့
        "prices": {"1m": 35000, "3m": 87500, "12m": 280000}
    }
}

def get_plan_price(plan, duration="1m"):
    return PLAN_LIMITS[plan]["prices"].get(duration, PLAN_LIMITS[plan]["price"])

def get_duration_label(duration):
    labels = {"1m": "၁ လ", "3m": "၃ လ", "12m": "၁၂ လ"}
    return labels.get(duration, duration)

def get_price_usd(price_mmk):
    return round(price_mmk / EXCHANGE_RATE, 2)

BOT_NAMES = ["မစ္စတာသန်း", "ကိုသန်း", "သန်း"]

# ====== Flask App ======
flask_app = Flask(__name__)

def get_db_connection():
    try:
        return sqlite3.connect("bot_users.db", check_same_thread=False)
    except Exception as e:
        logger.error(f"❌ DB connection error: {e}")
        raise

def check_auth(username, password):
    return username in ADMIN_USERS and ADMIN_USERS[username] == password

def authenticate():
    return Response(
        "❌ Unauthorized! Username and Password required.",
        401,
        {"WWW-Authenticate": 'Basic realm="Login Required"'},
    )

def requires_auth(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

# ====== Flask Routes ======
@flask_app.route("/")
def home():
    return "🤖 Bot is running! Visit /admin/proofs for dashboard."

@flask_app.route("/health")
def health():
    return "OK", 200

@flask_app.route("/admin/proofs")
@requires_auth
def admin_proofs():
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT user_id, plan, usage_count, proof_status, proof_file_id FROM users")
    results = c.fetchall()
    conn.close()

    html = "<h2>📋 Proofs Dashboard</h2>"
    html += "<table border='1' cellpadding='5' style='border-collapse:collapse;'>"
    html += "<tr><th>User ID</th><th>Plan</th><th>Usage</th><th>Proof Status</th><th>File ID</th><th>Actions</th></tr>"
    for uid, plan, usage, status, file_id in results:
        html += (
            f"<tr><td>{uid}</td><td>{plan}</td><td>{usage}</td><td>{status}</td>"
            f"<td>{file_id[:30] if file_id else '-'}...</td>"
        )
        if status == "pending":
            html += (
                f"<td><a href='/admin/approve/{uid}'>✅ Approve</a> | "
                f"<a href='/admin/reject/{uid}'>❌ Reject</a></td>"
            )
        else:
            html += "<td>-</td>"
        html += "</tr>"
    html += "</table>"
    return html

@flask_app.route("/admin/users")
@requires_auth
def admin_users():
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT user_id, plan, usage_count, proof_status FROM users ORDER BY user_id")
    results = c.fetchall()
    conn.close()

    html = "<h2>👥 User List</h2>"
    html += "<table border='1' cellpadding='5' style='border-collapse:collapse;'>"
    html += "<tr><th>User ID</th><th>Plan</th><th>Usage</th><th>Proof Status</th></tr>"
    for uid, plan, usage, status in results:
        html += f"<tr><td>{uid}</td><td>{plan}</td><td>{usage}</td><td>{status}</td></tr>"
    html += "</table>"
    html += f"<br><p><b>Total Users: {len(results)}</b></p>"
    return html

@flask_app.route("/admin/stats")
@requires_auth
def admin_stats():
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM users")
    total_users = c.fetchone()[0]
    c.execute("SELECT plan, COUNT(*) FROM users GROUP BY plan")
    plan_stats = c.fetchall()
    c.execute("SELECT COUNT(*) FROM users WHERE proof_status='pending'")
    pending = c.fetchone()[0]
    c.execute("SELECT SUM(usage_count) FROM users")
    total_usage = c.fetchone()[0] or 0
    conn.close()

    plan_names = [row[0] for row in plan_stats]
    plan_counts = [row[1] for row in plan_stats]

    html = f"""
    <h2>📊 Bot Statistics</h2>
    <p><b>Total Users:</b> {total_users}</p>
    <p><b>Pending Proofs:</b> {pending}</p>
    <p><b>Total API Calls:</b> {total_usage}</p>
    
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
    <div style="width: 500px; height: 350px; margin: 20px auto;">
        <canvas id="planChart"></canvas>
    </div>

    <script>
        const ctx = document.getElementById('planChart').getContext('2d');
        const chart = new Chart(ctx, {{
            type: 'bar',
            data: {{
                labels: {json.dumps(plan_names)},
                datasets: [{{
                    label: 'အသုံးပြုသူ အရေအတွက်',
                    data: {json.dumps(plan_counts)},
                    backgroundColor: 'rgba(54, 162, 235, 0.6)',
                    borderColor: 'rgba(54, 162, 235, 1)',
                    borderWidth: 1
                }}]
            }},
            options: {{ scales: {{ y: {{ beginAtZero: true }} }} }}
        }});
    </script>
    """
    return html

@flask_app.route("/download_db")
@requires_auth
def download_db():
    try:
        return send_file("bot_users.db", as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

@flask_app.route("/admin/approve/<user_id>")
@requires_auth
def approve_user(user_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT plan FROM users WHERE user_id=? AND proof_status='pending'", (user_id,))
    result = c.fetchone()
    if not result:
        conn.close()
        return f"❌ User {user_id} not found or not pending."
    plan = result[0]
    price = PLAN_LIMITS[plan]["price"]
    c.execute(
        "UPDATE users SET proof_status='approved', usage_count=0, price=? WHERE user_id=? AND proof_status='pending'",
        (price, user_id),
    )
    conn.commit()
    conn.close()
    return f"✅ User {user_id} upgraded to {plan} Plan!"

@flask_app.route("/admin/reject/<user_id>")
@requires_auth
def reject_user(user_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute(
        "UPDATE users SET proof_status='rejected' WHERE user_id=? AND proof_status='pending'",
        (user_id,),
    )
    if c.rowcount == 0:
        conn.close()
        return f"❌ User {user_id} not found or not pending."
    conn.commit()
    conn.close()
    return f"❌ User {user_id} proof rejected!"

# ====== Utility Functions ======
def get_bot_name(text):
    for name in BOT_NAMES:
        if name.lower() in text.lower():
            return name
    return None

def clean_text(text):
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"t\.me/\S+", "", text)
    text = re.sub(r"www\.\S+", "", text)
    text = re.sub(r"\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def calculate_expiry(duration):
    """သက်တမ်းအလိုက် Expiry Date ကို တွက်ချက်မယ်"""
    if duration == "1m":
        delta = timedelta(days=30)
    elif duration == "3m":
        delta = timedelta(days=90)
    elif duration == "12m":
        delta = timedelta(days=365)
    else:
        delta = timedelta(days=30)
    return (datetime.utcnow() + delta).isoformat()

def get_plan_price(plan, duration="1m"):
    """သက်တမ်းအလိုက် ဈေးနှုန်းကို ပြန်ပေးမယ်"""
    return PLAN_LIMITS[plan]["prices"].get(duration, PLAN_LIMITS[plan]["price"])

def get_duration_label(duration):
    labels = {"1m": "၁ လ", "3m": "၃ လ", "12m": "၁၂ လ"}
    return labels.get(duration, duration)


# ====== Database Backup Function ======
def backup_and_send(bot):
    try:
        conn = sqlite3.connect("bot_users.db", check_same_thread=False)
        with open("bot_users_backup.db", "w", encoding='utf-8') as f:
            for line in conn.iterdump():
                f.write(f"{line}\n")
        conn.close()
        
        with open("bot_users_backup.db", "rb") as f:
            bot.send_document(
                chat_id=ADMIN_IDS[0],
                document=f,
                caption=f"📦 Database Backup - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
        
        os.remove("bot_users_backup.db")
        logger.info("✅ Database backup sent to admin")
    except Exception as e:
        logger.error(f"❌ Backup error: {e}")

# ====== Database Functions ======
def init_db():
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    
    # ✅ Existing tables (unchanged)
    for col in ["birthdate"]:
        try:
            c.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT")
        except sqlite3.OperationalError:
            pass
    
            c.execute("""CREATE TABLE IF NOT EXISTS users (
        user_id TEXT PRIMARY KEY,
        plan TEXT DEFAULT 'free',
        usage_count INTEGER DEFAULT 0,
        proof_status TEXT DEFAULT 'none',
        proof_file_id TEXT,
        price INTEGER DEFAULT 0,
        proof_timestamp TEXT,
        goals TEXT,
        weaknesses TEXT,
        dream TEXT,
        career TEXT,
        money_mindset TEXT,
        relationship TEXT,
        birthdate TEXT,
        duration TEXT DEFAULT '1m',
        expiry_date TEXT
    )""")
    
    # ✅ Migration
    for col in ["duration", "expiry_date"]:
        try:
            c.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT")
        except sqlite3.OperationalError:
            pass
    
            c.execute("""CREATE TABLE IF NOT EXISTS referrals (
        inviter_id TEXT, invited_id TEXT, timestamp TEXT
    )""")
    
    c.execute("""CREATE TABLE IF NOT EXISTS habits (
        user_id TEXT, habit TEXT, created_at TEXT
    )""")
    
    c.execute("""CREATE TABLE IF NOT EXISTS response_cache (
        query TEXT PRIMARY KEY,
        response TEXT,
        created_at TEXT
    )""")
    
    # ✅ NEW: Transactions table for auto-verify
    c.execute("""CREATE TABLE IF NOT EXISTS transactions (
        tx_id TEXT PRIMARY KEY,
        user_id TEXT,
        plan TEXT,
        timestamp TEXT,
        used BOOLEAN DEFAULT 0,
        expiry_date TEXT
    )""")
    
    c.execute("CREATE INDEX IF NOT EXISTS idx_transactions_user ON transactions(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_transactions_plan ON transactions(plan)")
    
    # ✅ Migration for existing columns
    for col in ["proof_status", "proof_file_id", "price", "proof_timestamp", "goals", "weaknesses", "dream", "career", "money_mindset", "relationship", "birthdate"]:
        try:
            c.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT")
        except sqlite3.OperationalError:
            pass
    
    conn.commit()
    conn.close()
    logger.info("✅ Database initialized successfully.")

def add_user(user_id, plan="free", duration="1m"):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    price = get_plan_price(plan, duration)
    expiry_date = calculate_expiry(duration)
    c.execute("""INSERT OR REPLACE INTO users (
        user_id, plan, usage_count, proof_status, proof_file_id, price, proof_timestamp,
        goals, weaknesses, dream, career, money_mindset, relationship, birthdate,
        duration, expiry_date
    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
    (user_id, plan, 0, "none", None, price, None, None, None, None, None, None, None, None, duration, expiry_date))
    conn.commit()
    conn.close()

def get_user(user_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT plan, usage_count, proof_status, proof_file_id, price, proof_timestamp, goals, weaknesses, dream, career, money_mindset, relationship, birthdate FROM users WHERE user_id=?", (user_id,))
    result = c.fetchone()
    conn.close()
    return result

def normalize_text(text: str) -> str:
    return re.sub(r'\s+', ' ', text.strip()).lower()

def get_cached_response(query: str) -> str | None:
    cache_key = normalize_text(query)
    try:
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("SELECT response, created_at FROM response_cache WHERE query=?", (cache_key,))
        result = c.fetchone()
        conn.close()
        if result:
            response, created_at = result
            if (datetime.utcnow() - datetime.fromisoformat(created_at)).days <= 7:
                return response
        return None
    except Exception as e:
        logger.error(f"❌ Get cache error: {e}")
        return None

def save_cached_response(query: str, response: str):
    cache_key = normalize_text(query)
    try:
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("INSERT OR REPLACE INTO response_cache (query, response, created_at) VALUES (?, ?, ?)",
                  (cache_key, response, datetime.utcnow().isoformat()))
        conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"❌ Save cache error: {e}")

def update_profile(user_id, field, value):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute(f"UPDATE users SET {field}=? WHERE user_id=?", (value, user_id))
    conn.commit()
    conn.close()

def update_proof(user_id, file_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("UPDATE users SET proof_file_id=?, proof_status='pending', proof_timestamp=? WHERE user_id=?", (file_id, datetime.utcnow().isoformat(), user_id))
    conn.commit()
    conn.close()

def check_limit(user_id):
    user = get_user(user_id)
    if not user:
        add_user(user_id, "free")
        return True
    plan = user[0]
    usage = user[1]
    if usage >= PLAN_LIMITS[plan]["limit"]:
        return False
    return True

def increment_usage(user_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("UPDATE users SET usage_count = usage_count + 1 WHERE user_id=?", (user_id,))
    conn.commit()
    conn.close()

# ====== Referral System ======
def generate_ref_code(user_id):
    return f"REF{user_id}"

def give_referral_reward(inviter_id, invited_id):
    try:
        conn = sqlite3.connect("bot_users.db", check_same_thread=False)
        c = conn.cursor()
        today = datetime.utcnow().date().isoformat()
        c.execute("SELECT COUNT(*) FROM referrals WHERE inviter_id=? AND timestamp LIKE ?", (inviter_id, f"{today}%"))
        daily_count = c.fetchone()[0]
        if daily_count >= 3:
            logger.warning(f"⚠️ Referral limit reached for {inviter_id}")
            conn.close()
            return
        c.execute("SELECT usage_count FROM users WHERE user_id=?", (inviter_id,))
        row = c.fetchone()
        if not row:
            add_user(inviter_id, "free")
            usage = 0
        else:
            usage = row[0]
        new_usage = max(usage - 50, 0)
        c.execute("UPDATE users SET usage_count=? WHERE user_id=?", (new_usage, inviter_id))
        c.execute("INSERT INTO referrals (inviter_id, invited_id, timestamp) VALUES (?, ?, ?)", (inviter_id, invited_id, datetime.utcnow().isoformat()))
        conn.commit()
        conn.close()
        logger.info(f"🎁 Referral reward (50 calls): inviter={inviter_id}, invited={invited_id}")
    except Exception as e:
        logger.error(f"❌ Referral reward error: {e}")

def get_referral_count(user_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM referrals WHERE inviter_id=?", (user_id,))
    count = c.fetchone()[0]
    conn.close()
    return count

# ====== Habits ======
def add_habit(user_id, habit_text):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("INSERT INTO habits (user_id, habit, created_at) VALUES (?, ?, ?)", (user_id, habit_text, datetime.utcnow().isoformat()))
    conn.commit()
    conn.close()

def get_habits(user_id):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT habit, created_at FROM habits WHERE user_id=? ORDER BY created_at DESC LIMIT 20", (user_id,))
    results = c.fetchall()
    conn.close()
    return results

def reset_usage():
    try:
        conn = sqlite3.connect("bot_users.db", check_same_thread=False)
        c = conn.cursor()
        c.execute("UPDATE users SET usage_count = 0")
        conn.commit()
        conn.close()
        logger.info("✅ Monthly usage reset completed successfully.")
    except Exception as e:
        logger.error(f"❌ Usage reset error: {e}")

# ====== Response Cache ======
def save_cached_response(query: str, response: str):
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute(
        "INSERT OR REPLACE INTO response_cache (query, response, created_at) VALUES (?, ?, ?)",
        (query, response, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()

# ====== Fuzzy Matching Helpers ======
def normalize_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^\u1000-\u109F0-9a-zA-Z\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def fuzzy_score(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()

# ====== Local Responses ======
LOCAL_RESPONSES = {
    "hello": "ဟယ်လို! မင်္ဂလာပါဗျ။ ဘာကူညီပေးရမလဲ?",
    "hi": "ဟိုင်း! ဒီနေ့ ဘာတွေ လုပ်နေလဲဗျ။",
    "ဟိုင်း": "ဟိုင်း! ဘာမေးချင်လဲဗျ။",
    "မင်္ဂလာပါ": "မင်္ဂလာပါဗျ။ ကြိုဆိုပါတယ်။",
    "နေကောင်းလား": "ကျေးဇူးတင်ပါတယ်၊ ကျွန်တော်ကတော့ အဆင်ပြေပါတယ်။ ခင်ဗျားကရော?",
    "ဘာလုပ်နေလဲ": "ခင်ဗျားကို ကူညီဖို့ စောင့်နေတာပေါ့ဗျာ။",
    "အိမ်မှာလား": "ဟုတ်ကဲ့ဗျ၊ ဒီမှာပါပဲ။ ဘာမေးချင်လဲဗျ။",
    "ဘယ်လိုလဲ": "ကျွန်တော်ကတော့ အဆင်ပြေပါတယ်ဗျ။ ခင်ဗျားကရော?",
    "စားပြီးပြီလား": "ရပါတယ်ဗျ။ မင်းစားပြီးပြီလား?",
    "thanks": "ရပါတယ်ဗျ။ ကျေးဇူးတင်ပါတယ်။",
    "thank you": "ရပါတယ်ဗျ။ ကျေးဇူးတင်ပါတယ်။",
    "ကျေးဇူးတင်ပါတယ်": "အေးပါဗျ။ အားမနာပါနဲ့။",
    "ကျေးဇူး": "အေးပါဗျ။ ကျေးဇူးတင်ပါတယ်။",
    "အဆင်ပြေပါတယ်": "ဝမ်းသာပါတယ်ဗျ။",
    "အင်း": "အင်း ဟုတ်ကဲ့ဗျ။",
    "ဟုတ်": "ဟုတ်ပါတယ်ဗျ။",
    "မဟုတ်ဘူး": "မဟုတ်ဘူးဆိုတော့ ဘာဖြစ်လဲဗျ။",
    "အိုကေ": "အိုကေပါဗျ။",
    "ok": "Okay ပါဗျ။",
    "ကောင်းပြီ": "ကောင်းပါပြီဗျ။",
    "သိပြီ": "သိရင် ဝမ်းသာပါတယ်ဗျ။",
    "မသိဘူး": "မသိရင် ဘာကို သိချင်တာလဲဗျ။",
    "ရပါတယ်": "ရပါတယ်ဗျ။",
    "ဘယ်သူလဲ": "ကျွန်တော်က မစ္စတာသန်း (Mr.T) ပါ။ ခင်ဗျားရဲ့ ကိုယ်ပိုင်အကြံပေးဘော့ပါ။",
    "မင်းကဘာလဲ": "ကျွန်တော်က မစ္စတာသန်း (Mr.T) ပါ။ ခင်ဗျားရဲ့ ကိုယ်ပိုင်အကြံပေးဘော့ပါ။",
    "ဘာလုပ်ပေးနိုင်လဲ": "လုပ်ငန်း၊ နည်းပညာ၊ စိုက်ပျိုးရေး၊ ဗီဒီယို၊ စီးပွားရေး အပါအဝင် အကြံဉာဏ် (၁၆) မျိုး ပေးနိုင်ပါတယ်ဗျ။",
    "ဘာတွေလုပ်ပေးနိုင်လဲ": "ဘဝအကြံဉာဏ်၊ စီးပွားရေး၊ နည်းပညာ၊ ကျန်းမာရေး စသဖြင့် အကုန်ကူညီနိုင်ပါတယ်။",
    "help": "ကူညီမှုအတွက် /help ကို နှိပ်ပါ။",
    "အကူအညီ": "ဘာကူညီပေးရမလဲဗျ။ ဥပမာ - /ask နည်းပညာအကြောင်း",
    "အမိန့်": "Command တွေကို /help မှာ ကြည့်နိုင်ပါတယ်ဗျ။",
    "command": "Command တွေကို /help မှာ ကြည့်နိုင်ပါတယ်ဗျ။",
    "စတင်မယ်": "စတင်လိုက်ပါပြီဗျ။ ဘာမေးမလဲ။",
    "price": "Plan ဈေးနှုန်းတွေကို /subscribe နှိပ်ပြီး ကြည့်နိုင်ပါတယ်ဗျ။",
    "စျေးနှုန်း": "စျေးနှုန်းတွေ သိချင်ရင် /subscribe နှိပ်ပါဗျ။",
    "plan": "Plan တွေက Free, Basic, Premium နဲ့ Premium+ ရှိပါတယ်။ /subscribe နှိပ်ပါ။",
    "ပလန်": "Plan တွေက Free, Basic, Premium နဲ့ Premium+ ရှိပါတယ်။ /subscribe နှိပ်ပါ။",
    "ဘယ်လိုဝယ်ရမလဲ": "Plan ဝယ်ဖို့ဆို /subscribe နှိပ်ပြီး ရွေးပါ၊ ပြီးရင် proof ပို့ပါ။",
    "ငွေလွှဲ": "ငွေလွှဲရန်: KBZPay: 09426419462 | WavePay: 09426419462 ။ ပြီးရင် proof ပို့ပါ။",
    "ပိုက်ဆံ": "ပိုက်ဆံပေးချေမှုအတွက် /subscribe ကို နှိပ်ပါဗျ။",
    "subscribe": "Plan ရွေးရန် အောက်ပါခလုတ်များကို နှိပ်ပါ။",
    "ဝယ်မယ်": "ဝယ်ချင်ရင် /subscribe နှိပ်ပြီး Plan ရွေးပါ။",
    "proof": "ငွေလွှဲပြီးရင် Screenshot ကို ဒီထဲ Photo ပို့ပါ။",
    "ပြေစာ": "ပြေစာဓာတ်ပုံကို ပို့ပေးပါ။",
    "ဗီဒီယိုဖန်တီးနည်း": "ဗီဒီယိုဖန်တီးနည်းအသေးစိတ်၊ Software အကြံပြုချက်တွေကို /ask နဲ့ မေးကြည့်နိုင်ပါတယ်ဗျ။",
    "video editing": "Video editing အကြောင်း သိချင်ရင် /ask မှာ မေးပါ။",
    "capcut": "CapCut အကြောင်း သိချင်ရင် /ask မှာ မေးပါ။",
    "ဒီဇိုင်းဆွဲနည်း": "ဒီဇိုင်းဆွဲနည်းအတွက် /ask မှာ မေးပါဗျ။",
    "စိုက်ပျိုးနည်း": "စိုက်ပျိုးနည်းအသေးစိတ်ကို /ask နဲ့ မေးပါ။",
    "ကြက်မွေးနည်း": "ကြက်မွေးနည်းအသေးစိတ်ကို /ask နဲ့ မေးပါ။",
    "ကွန်ပျူတာပြဿနာ": "ကွန်ပျူတာပြဿနာရှိရင် /ask နဲ့ မေးကြည့်ပါဗျ။",
    "နည်းပညာ": "နည်းပညာအကြံဉာဏ်တွေ လိုချင်ရင် /ask နဲ့ မေးပါ။",
    "ဗေဒင်": "ဗေဒင်မေးချင်ရင် မင်းရဲ့ မွေးနေ့ကို /profile birthdate : <မွေးနေ့> လို့ သိမ်းထားပါ၊ ပြီးရင် /ask မှာ မေးပါ။",
    "ရာသီခွင်": "ရာသီခွင်အကြောင်း (ဥပမာ - မိဿ၊ ပြိဿ စသည်) သိချင်ရင် /ask မှာ မေးပါဗျ။",
    "နက္ခတ်ဗေဒင်": "နက္ခတ်ဗေဒင်အသေးစိတ်ကို /ask မှာ မေးပါ။",
    "မွေးနေ့": "မွေးနေ့နဲ့ ကံကြမ္မာသိချင်ရင် /profile birthdate : <သင့်မွေးနေ့> ထည့်ပြီး /ask မှာ မေးပါ။",
    "ကံကြမ္မာ": "ကံကြမ္မာဖတ်ချင်ရင် မွေးနေ့လိုအပ်လို့ /profile birthdate : <သင့်မွေးနေ့> ထည့်ပါ။",
    "အတိတ်နိမိတ်": "အတိတ်နိမိတ် (ဥပမာ - ကြောင်နှာချေခြင်း၊ ငှက်မြည်ခြင်း) အကြောင်း သိချင်ရင် /ask မှာ မေးပါ။",
    "သိုင်းပညာ": "သိုင်းပညာ၊ ကိုယ်ခံပညာ၊ လက်ဝှေ့အကြောင်း အသေးစိတ်သိချင်ရင် /ask မှာ မေးပါဗျ။",
    "လက်ဝှေ့": "လက်ဝှေ့အနုပညာအကြောင်း အသေးစိတ်ကို /ask မှာ မေးပါ။",
    "ကိုယ်ခံပညာ": "ကိုယ်ခံပညာသင်ဖို့ ဘယ်လိုစမလဲဆိုတာ /ask မှာ မေးပါ။",
    "သိုင်း": "သိုင်းပညာနဲ့ ပတ်သက်ပြီး ဘာသိချင်လဲဗျ။ /ask မှာ မေးကြည့်ပါ။",
    "ကရာတေး": "ကရာတေးအကြောင်း သိချင်ရင် /ask မှာ မေးပါ။",
    "တိုက်ခိုက်နည်း": "တိုက်ခိုက်ရေး၊ ခုခံရေး နည်းပညာတွေအတွက် /ask မှာ မေးပါ။",
    # ဗေဒင် (Astrology)
    "ဗေဒင်": "ဗေဒင်မေးချင်ရင် မင်းရဲ့ မွေးနေ့ကို /profile birthdate : <မွေးနေ့> လို့ သိမ်းထားပါ၊ ပြီးရင် /ask မှာ မေးပါ။",
    "ရာသီခွင်": "ရာသီခွင်အကြောင်း (ဥပမာ - မိဿ၊ ပြိဿ စသည်) သိချင်ရင် /ask မှာ မေးပါဗျ။",
    "နက္ခတ်ဗေဒင်": "နက္ခတ်ဗေဒင်အသေးစိတ်ကို /ask မှာ မေးပါ။",
    "မွေးနေ့": "မွေးနေ့နဲ့ ကံကြမ္မာသိချင်ရင် /profile birthdate : <သင့်မွေးနေ့> ထည့်ပြီး /ask မှာ မေးပါ။",
    "ကံကြမ္မာ": "ကံကြမ္မာဖတ်ချင်ရင် မွေးနေ့လိုအပ်လို့ /profile birthdate : <သင့်မွေးနေ့> ထည့်ပါ။",
    "အတိတ်နိမိတ်": "အတိတ်နိမိတ် (ဥပမာ - ကြောင်နှာချေခြင်း၊ ငှက်မြည်ခြင်း) အကြောင်း သိချင်ရင် /ask မှာ မေးပါ။",

    # ရာသီဥတု (Weather)
    "ရာသီဥတု": "ရာသီဥတုအခြေအနေကို ခန့်မှန်းဖို့ /ask မှာ သင့်မြို့နယ်ကို ထည့်ပြီး မေးပါ။",
    "မိုးရွာမလား": "မိုးရွာနိုင်ခြေကို /ask မှာ သင့်တည်နေရာထည့်ပြီး မေးကြည့်ပါဗျ။",
    
    # ရင်းနှီးမြှုပ်နှံမှု (Investment)
    "ရင်းနှီးမြှုပ်နှံ": "ရင်းနှီးမြှုပ်နှံမှုအကြောင်း အသေးစိတ် (စတော့၊ ရွှေ၊ အိမ်ခြံမြေ) သိချင်ရင် /ask မှာ မေးပါဗျ။",
    "စတော့": "စတော့ဈေးကွက်အကြောင်း တစ်ကိုယ်ရေအကြံဉာဏ်အတွက် /ask မှာ မေးပါ။",
    
    # ခရစ်ပတို (Crypto)
    "ခရစ်ပတို": "ခရစ်ပတိုငွေကြေးတွေအကြောင်း အန္တရာယ်ကင်းစွာ လေ့လာချင်ရင် /ask မှာ မေးပါ။",
    "bitcoin": "Bitcoin အကြောင်း အသေးစိတ်ကို /ask မှာ မေးပါ။",
    
    # အော်ဒီယိုစာအုပ် (Audiobook)
    "အော်ဒီယို": "နားထောင်သင့်တဲ့ စာအုပ်တွေအတွက် /ask မှာ မေးပါ။",
    "စာအုပ်": "စာအုပ်အကြံပြုချက်တွေအတွက် /ask မှာ မေးပါ။",
    
    # ပညာရေး (Education)
    "ပညာရေး": "ပညာရေးဆိုင်ရာ လမ်းညွှန်ချက်တွေအတွက် /ask မှာ မေးပါ။",
    "ကျောင်း": "ကျောင်းရွေးချယ်ခြင်းနှင့် သင်တန်းများအတွက် /ask မှာ မေးပါ။",
    
    # ဘူမိဗေဒ (Geology)
    "ဘူမိ": "မြေအနေအထားနှင့် ကျောက်တုံးများအကြောင်း /ask မှာ မေးပါ။",
    "မြေငလျင်": "ငလျင်အန္တရာယ်အတွက် /ask မှာ မေးပါ။",
    
    # ဆေးဝါး (Pharmacist)
    "ဆေးဝါး": "ဆေးဝါးအသုံးပြုနည်းနဲ့ ပတ်သက်လို့ အခြေခံသိချင်ရင် /ask မှာ မေးပါ။ (ဆရာဝန်နဲ့ တိုင်ပင်ဖို့ မမေ့ပါနဲ့)",
    "ဆေး": "ဆေးဝါးအကြောင်း အသေးစိတ်သိချင်ရင် /ask မှာ မေးပါ။",
    
    # အချစ်ရေး အိမ်ထောင်ရေး (Love & Marriage)
    "အချစ်ရေး": "အချစ်ရေးအကြံဉာဏ်တွေအတွက် /ask မှာ မေးပါ။",
    "အိမ်ထောင်": "အိမ်ထောင်ရေးပြဿနာများအတွက် စိတ်ရှည်ရှည်နဲ့ /ask မှာ မေးပါ။",
    
    # ကလေးနာမည်ပေး (Baby Names)
    "ကလေးနာမည်": "ကလေးနာမည်ကောင်းများ ရွေးချင်ရင် /profile birthdate : <မွေးနေ့> ထည့်ပြီး /ask မှာ မေးပါ။",
    "နာမည်ပေး": "ကလေးနာမည်ပေးဖို့ မွေးနေ့လိုအပ်လို့ /profile birthdate : <မွေးနေ့> ထည့်ပါ။",
    
    # နေ့စဉ်သုံးစကားများ (Daily Phrases)
    "နေ့စဉ်သုံး": "နေ့စဉ်သုံးအင်္ဂလိပ်စကားပြောများ လေ့လာချင်ရင် /ask မှာ မေးပါ။",
    "အင်္ဂလိပ်စကား": "အင်္ဂလိပ်စကားပြော အသုံးအနှုန်းတွေအတွက် /ask မှာ မေးပါ။",
    
    # အလာဘ သလာဘ (Small Talk)
    "အလာဘ": "စကားစမြည်ပြောနည်းနဲ့ မိတ်ဆက်စကားတွေအတွက် /ask မှာ မေးပါ။",
    "သလာဘ": "အလာဘသလာဘစကားများနှင့် အသုံးအနှုန်းများအတွက် /ask မှာ မေးပါ။",
    "ထိုင်းထီ": "ထိုင်းထီ (ချဲထီ) အကြောင်း အသေးစိတ်သိချင်ရင် /ask မှာ မေးပါဗျ။ ဒါပေမယ့် ဒါဟာ မြန်မာပြည်မှာ တရားမဝင်တဲ့ လောင်းကစားတစ်ခုဖြစ်ပြီး အန္တရာယ်ရှိနိုင်ကြောင်း သတိပေးချင်ပါတယ်။",
    "ချဲထီ": "ချဲထီ ဆိုတာ ထိုင်းအစိုးရထီကို မြန်မာပြည်မှာ ခေတ်စားတဲ့အခေါ်ပါ။ ထီပုံစံနဲ့ ကံစမ်းနည်းတွေကို /ask မှာ မေးကြည့်နိုင်ပါတယ်။",
    "နှစ်လုံးထီ": "နှစ်လုံးထီ ဟာ ထိုင်း SET စတော့ဈေးကွက်ပိတ်ချိန် ကိန်းဂဏန်းတွေကို အခြေခံပြီး ထိုးရတဲ့ပုံစံပါ [citation:1]။ အသေးစိတ်ကို /ask မှာ မေးပါ။",
    "သုံးလုံးထီ": "သုံးလုံးထီ ဟာ ထိုင်းအစိုးရထီရဲ့ နံပါတ်နောက်ဆုံး (၃) လုံးကို အခြေခံပြီး ထိုးတာပါ။ အသေးစိတ်ကို /ask မှာ မေးပါ။",
    "ထီနံပါတ်": "ထီနံပါတ်ရွေးချယ်နည်း အကြံဉာဏ်တွေအတွက် /ask မှာ မေးပါ။",
    # ကားအရောင်းအဝယ် (Car Trading)
    "ကားအရောင်း": "ကားအရောင်းအဝယ်ပြုလုပ်ခြင်း၊ ဈေးနှုန်းနှင့် စာရွက်စာတမ်းများအတွက် /ask မှာ မေးပါဗျ။",
    "ကားဝယ်": "ကားဝယ်ယူရာတွင် စစ်ဆေးသင့်သည့်အချက်များအတွက် /ask မှာ မေးပါ။",
    
    # ကွန်ပျူတာအရောင်းအဝယ် (Computer Trading)
    "ကွန်ပျူတာအရောင်း": "ကွန်ပျူတာ၊ လက်ပ်တော့ အရောင်းအဝယ်အတွက် /ask မှာ မေးပါဗျ။",
    "လက်ပ်တော့": "လက်ပ်တော့ရွေးချယ်ခြင်းနှင့် အရောင်းအဝယ်အတွက် /ask မှာ မေးပါ။",
    
    # ကျောက်မျက်ရတနာ (Gemstone)
    "ကျောက်မျက်ရတနာ": "ကျောက်မျက်ရတနာအမျိုးအစားနှင့် အရည်အသွေးစစ်ဆေးနည်းအတွက် /ask မှာ မေးပါ။",
    "ကျောက်": "ကျောက်မျက်ရတနာရွေးချယ်နည်းအတွက် /ask မှာ မေးပါ။",
    
    # အဝတ်အထည် (Clothing)
    "အဝတ်အထည်": "အဝတ်အထည်အရောင်းအဝယ်နှင့် ဖက်ရှင်ရွေးချယ်နည်းအတွက် /ask မှာ မေးပါ။",
    "ဖက်ရှင်": "ဖက်ရှင်စတိုင်များနှင့် အဝတ်အထည်ရွေးချယ်နည်းအတွက် /ask မှာ မေးပါ။",
    
    # ဖုန်းပြင် (Phone Repair)
    "ဖုန်းပြင်": "ဖုန်းပြုပြင်ခြင်းဆိုင်ရာ အကြံဉာဏ်တွေအတွက် /ask မှာ မေးပါ။",
    "ဖုန်းရေစို": "ဖုန်းရေစိုပြဿနာအတွက် ချက်ချင်းလုပ်ဆောင်သင့်သည်များကို /ask မှာ မေးပါ။",
    
    # ရွှေအရောင်းအဝယ် (Gold Trading)
    "ရွှေအရောင်း": "ရွှေဈေးနှုန်းနှင့် အရည်အသွေးစစ်ဆေးနည်းအတွက် /ask မှာ မေးပါ။",
    "ရွှေ": "ရွှေဝယ်ယူရောင်းချခြင်းနှင့် ရင်းနှီးမြှုပ်နှံမှုအတွက် /ask မှာ မေးပါ။",
        # အထည်ချုပ် (Tailoring)
    "အထည်ချုပ်": "အထည်ချုပ်လုပ်ငန်းနှင့် အဝတ်အထည်ချုပ်နည်းအတွက် /ask မှာ မေးပါဗျ။",
    "ချုပ်": "အဝတ်အထည်ချုပ်နည်းနှင့် အတိုင်းအတာယူနည်းများအတွက် /ask မှာ မေးပါ။",
    
    # ဒီဇိုင်း (Design)
    "ဒီဇိုင်နာ": "ဒီဇိုင်းပညာရပ်များအတွက် /ask မှာ မေးပါ။",
    "ဂရပ်ဖစ်": "ဂရပ်ဖစ်ဒီဇိုင်းနှင့် အနုပညာဖန်တီးနည်းများအတွက် /ask မှာ မေးပါ။",
    
    # ဓာတု (Chemical)
    "ဓာတု": "ကုန်ထုတ်ဓာတုပညာရပ်များအတွက် /ask မှာ မေးပါ။",
    
    # ရေမွှေး/ဆပ်ပြာ
    "ရေမွှေး": "ရေမွှေးနှင့် အလှကုန်ပစ္စည်းထုတ်လုပ်နည်းအတွက် /ask မှာ မေးပါ။",
    "ဆပ်ပြာ": "ဆပ်ပြာနှင့် သန့်ရှင်းရေးပစ္စည်းထုတ်လုပ်နည်းအတွက် /ask မှာ မေးပါ။",
    
    # ဘိုင်နရီ/ဖြူးချား
    "ဘိုင်နရီ": "ဘက်ထရီပြုပြင်ထိန်းသိမ်းနည်းအတွက် /ask မှာ မေးပါ။",
    "ဖြူးချား": "ဖျူးစ်/ဘရိတ်ကာနှင့် လျှပ်စစ်ပစ္စည်းများအကြောင်း /ask မှာ မေးပါ။",
    
    # ပတ်ဝန်းကျင်
    "ပတ်ဝန်းကျင်": "သဘာဝပတ်ဝန်းကျင်ထိန်းသိမ်းရေးအတွက် /ask မှာ မေးပါ။",
    "တိရိစ္ဆာန်": "တောရိုင်းတိရိစ္ဆာန်နှင့် သားငါးထိန်းသိမ်းရေးအတွက် /ask မှာ မေးပါ။",
    
    # ဘဏ် (Bank)
    "ဘဏ်": "ဘဏ်ဝန်ဆောင်မှုများနှင့် စာရင်းဖွင့်နည်းအတွက် /ask မှာ မေးပါ။",
    "ချေးငွေ": "ချေးငွေနှင့် အတိုးနှုန်းများအတွက် /ask မှာ မေးပါ။",
    
    # အရောင်း (Sales)
    "အရောင်းသမား": "အရောင်းပညာရပ်များနှင့် ဖောက်သည်စကားပြောနည်းအတွက် /ask မှာ မေးပါ။",
    
    # မုန့် (Bakery)
    "မုန့်": "မုန့်လုပ်ငန်းနှင့် မုန့်ဖုတ်နည်းအတွက် /ask မှာ မေးပါ။",
    "ကိတ်": "ကိတ်မုန့်ပြုလုပ်နည်းအတွက် /ask မှာ မေးပါ။",
    
    # စားသောက်ဆိုင် (Restaurant)
    "စားသောက်ဆိုင်": "စားသောက်ဆိုင်ဖွင့်နည်းနှင့် စီမံခန့်ခွဲနည်းအတွက် /ask မှာ မေးပါ။",
    
    # အလှပြင် (Beauty)
    "မိတ်ကပ်": "မိတ်ကပ်ဖျောက်နည်းနှင့် အလှပြင်နည်းအတွက် /ask မှာ မေးပါ။",
    "ဆံပင်": "ဆံပင်ညှပ်နည်းနှင့် ဆံပင်ဒီဇိုင်းအတွက် /ask မှာ မေးပါ။",
    
    # ဥယျာဉ် (Gardening)
    "ဥယျာဉ်": "ပန်းပျိုးခြင်းနှင့် ဥယျာဉ်စိုက်ပျိုးနည်းအတွက် /ask မှာ မေးပါ။",
    "ပန်း": "ပန်းစိုက်ပျိုးနည်းအတွက် /ask မှာ မေးပါ။",
    
    # တောင်တက် (Mountaineering)
    "တောင်တက်": "တောင်တက်ပညာနှင့် အန္တရာယ်ကင်းရှင်းရေးအတွက် /ask မှာ မေးပါ။",
    
    # မြေပုံ/ဒေသ (Geography)
    "မြေပုံ": "မြန်မာပြည်ဒေသမြေပုံနှင့် ခရီးသွားလမ်းညွှန်အတွက် /ask မှာ မေးပါ။",
    
    # ကမ္ဘာ့ရေးရာ (Global)
    "ကမ္ဘာ့": "ကမ္ဘာ့ရေးရာနှင့် နိုင်ငံတကာသတင်းများအတွက် /ask မှာ မေးပါ။",
    
    # နိုင်ငံရေး (Politics)
    "နိုင်ငံရေး": "နိုင်ငံရေးအခြေအနေများကို ဘက်မလိုက်ဘဲ သိရှိလိုပါက /ask မှာ မေးပါ။",
    
    # သတင်း (News)
    "သတင်း": "နောက်ဆုံးရသတင်းများအတွက် /ask မှာ မေးပါ။",
    
    # ပွဲစား (Broker)
    "ပွဲစား": "ကုန်သည်ပွဲစားလုပ်ငန်းနှင့် အဝယ်အရောင်းညှိနှိုင်းနည်းအတွက် /ask မှာ မေးပါ။",
    "ကုန်သည်": "ကုန်သည်လုပ်ငန်းနှင့် ဈေးနှုန်းအချက်အလက်များအတွက် /ask မှာ မေးပါ။",
    # Binary & Futures Trading
    "binary": "Binary Options သည် အနိုင်အရှုံးသာ ရှိပြီး မြန်မာပြည်မှာ တရားဝင်မဟုတ်ပါ။ အလွန်မြင့်မားတဲ့ Risk ရှိကြောင်း သိထားပါ။ /ask မှာ မေးနိုင်ပါတယ်။",
    "futures": "Futures Trading သည် Contract အပေါ်အခြေခံပြီး ဈေးကွက်အတက်အကျကို ခန့်မှန်းရတာပါ။ Risk များတဲ့အတွက် သေချာလေ့လာပြီးမှ /ask မှာ မေးကြည့်ပါ။",
    "binary option": "Binary Option သည် မြန်မာပြည်မှာ တရားဝင်မဟုတ်ပါ။ အန္တရာယ်ရှိနိုင်ကြောင်း သတိပြုပါ။ /ask မှာ အသေးစိတ်မေးနိုင်ပါတယ်။",
    "futures trading": "Futures Trading အကြောင်းအသေးစိတ်၊ Risk စီမံခန့်ခွဲနည်းများအတွက် /ask မှာ မေးပါ။",
    "leverage": "Leverage (အရင်းအနှီးချေးငွေ) သုံးပြီး ရောင်းဝယ်ခြင်းသည် အမြတ်ရရင်မြန်ပေမယ့် ဆုံးရှုံးရင်လည်း မြန်ပါတယ်။ ဒါကြောင့် /ask မှာ မေးပြီးမှ ဆုံးဖြတ်ပါ။",
        # သမိုင်း (History)
    "သမိုင်း": "မြန်မာ့သမိုင်း၊ ကမ္ဘာ့သမိုင်းအကြောင်း အသေးစိတ်သိချင်ရင် /ask မှာ မေးပါဗျ။",
    "ရှေးဟောင်း": "ရှေးဟောင်းသုတေသနဆိုင်ရာ အချက်အလက်များအတွက် /ask မှာ မေးပါ။",
    
    # ကုန်စျေးနှုန်း (Prices)
    "ကုန်စျေးနှုန်း": "ရွှေ၊ ဆန်၊ ဆီ အစရှိတဲ့ ကုန်စျေးနှုန်းတွေအတွက် /ask မှာ မေးပါ။",
    "စျေးနှုန်း": "ဒေသအလိုက် ကုန်စျေးနှုန်းအခြေအနေများအတွက် /ask မှာ မေးပါ။",
    
    # ရုပ်ရှင် (Movies)
    "ရုပ်ရှင်": "ရုပ်ရှင်ကောင်းများ အကြံပြုချက်ရယူရန် YouTube/Netflix လင့်နှင့်အတူ /ask မှာ မေးပါ။",
    "ကားရိုက်": "ရုပ်ရှင်ရိုက်ကူးနည်းပညာနှင့် ဇာတ်ကားအကြံပြုချက်များအတွက် /ask မှာ မေးပါ။",
    
    # စာပေ (Literature)
    "စာပေ": "ဝတ္ထု၊ ကဗျာရေးသားနည်းနှင့် စာပေဆိုင်ရာ အကြံဉာဏ်များအတွက် /ask မှာ မေးပါ။",
    "စာရေးဆရာ": "စာရေးဆရာများအကြောင်းနှင့် စာအုပ်အကြံပြုချက်များအတွက် /ask မှာ မေးပါ။",
    
    # ဘာသာရေး (Religion)
    "ဘာသာရေး": "ဘာသာတရားအမျိုးမျိုး၏ အခြေခံသဘောတရားများအတွက် /ask မှာ မေးပါ။",
    
    # ဝိပဿနာ (Vipassana)
    "ဝိပဿနာ": "တရားထိုင်နည်းနှင့် ဝိပဿနာကျင့်စဉ်များအတွက် /ask မှာ မေးပါ။",
    
    # ကျော်ဖြတ်နည်း (Overcoming)
    "အဆင့်နိမ့်": "ဆင်းရဲခြင်းမှ လွတ်မြောက်ရန် နည်းလမ်းများအတွက် /ask မှာ မေးပါ။",
    
    # ပန်းချီ/ပန်းပဲ/ပန်းထိမ်/ပန်းတမော့
    "ပန်းချီ": "ပန်းချီရေးဆွဲနည်းနှင့် ပညာရပ်များအတွက် /ask မှာ မေးပါ။",
    "ပန်းပဲ": "သံထည်ပစ္စည်းပြုလုပ်နည်းအတွက် /ask မှာ မေးပါ။",
    "ပန်းထိမ်": "လက်ဝတ်ရတနာပြုလုပ်နည်းအတွက် /ask မှာ မေးပါ။",
    "ပန်းတမော့": "ပန်းတမော့ပစ္စည်းပြုလုပ်နည်းအတွက် /ask မှာ မေးပါ။",
    
    # မြန်မာ့ပန်းဆယ်မျိုး (10 Arts)
    "ပန်းဆယ်မျိုး": "မြန်မာ့ရိုးရာအနုပညာဆယ်မျိုးအကြောင်း /ask မှာ မေးပါ။",
    
    # ဘတ်စ်ကား (Bus)
    "ဘတ်စ်ကား": "မြန်မာပြည်မြို့ကြီးများ၏ ဘတ်စ်ကားလမ်းကြောင်းများအတွက် /ask မှာ မေးပါ။",
    
    # ငွေရေးကြေးရေး (Financial)
    "ငွေရေးကြေးရေး": "မိသားစုဘဏ္ဍာရေးစီမံခန့်ခွဲနည်းများအတွက် /ask မှာ မေးပါ။",
    "ဘာသာပြန်": "ဘာသာစကားများ အပြန်အလှန် ဘာသာပြန်ဆိုခြင်းနှင့် သဒ္ဒါပြင်ဆင်ခြင်းအတွက် /ask မှာ မေးပါဗျ။",
    "ဘာသာစကား": "အင်္ဂလိပ်၊ တရုတ်၊ ထိုင်း စသည့် ဘာသာစကားများ သင်ယူလေ့လာလိုပါက /ask မှာ မေးပါ။",
    "အင်္ဂလိပ်စာ": "အင်္ဂလိပ်စကားပြော၊ သဒ္ဒါနှင့် ဘာသာပြန်ခြင်းအတွက် /ask မှာ မေးပါ။",
    "တရုတ်စာ": "တရုတ်စကားပြောနှင့် တရုတ်ဘာသာပြန်ခြင်းအတွက် /ask မှာ မေးပါ။",
}

    # ... (LOCAL_RESPONSES အကုန်လုံး အတိုင်းသား)


# ====== Optimized Local Response Function ======
def get_local_response(user_text: str) -> str | None:
    if not user_text:
        return None

    text_raw = user_text.strip()
    text_lower = text_raw.lower()
    if text_lower in LOCAL_RESPONSES:
        return LOCAL_RESPONSES[text_lower]

    norm = normalize_text(text_raw)
    if norm in LOCAL_RESPONSES:
        return LOCAL_RESPONSES[norm]

    best_key = None
    best_score = 0.0
    for key in LOCAL_RESPONSES.keys():
        score = fuzzy_score(norm, normalize_text(key))
        if score > best_score:
            best_score = score
            best_key = key

    if best_score >= 0.75 and best_key:
        return LOCAL_RESPONSES[best_key]

    return None

# ====== System Prompt ======
system_prompt = (
    "သင်ဟာ မစ္စတာသန်း (Mr.T) — funny, friendly, motivational AI Bot ဖြစ်ပါတယ်။\n"
    "မြန်မာပြည်က လူတွေအတွက် အကောင်းဆုံး ဘဝအကြံပေး၊ နည်းပညာရှင်၊ စီးပွားရေးလမ်းညွှန်နဲ့ နေ့စဉ်ပြဿနာဖြေရှင်းပေးသူ ဖြစ်ပါတယ်။\n\n"
    "မင်းရဲ့ ကျွမ်းကျင်မှု နယ်ပယ် (၁၂၀+) ခုကို အောက်ပါအတိုင်း ခွဲခြားထားပါတယ်။\n"
    "=== ဘဝနှင့် စိတ်ဓာတ် (Life & Motivation) ===\n"
    "1️⃣ Life Coach, 2️⃣ Motivation, 3️⃣ Mental Health & Stress Relief, 4️⃣ Motivational Monk Style (Dhamma-inspired, non-religious), 5️⃣ Overcome Low Environment, 6️⃣ Positive Thinking, 7️⃣ Emotional Intelligence, 8️⃣ Self-Confidence, 9️⃣ Time Management, 🔟 Productivity\n"
    "=== ငွေကြေးနှင့် စီးပွားရေး (Money & Business) ===\n"
    "1️⃣1️⃣ Money Mindset, 1️⃣2️⃣ Investment (Stock, Gold, Real Estate), 1️⃣3️⃣ Crypto, 1️⃣4️⃣ Binary & Futures Trading (⚠️ Risk Warning), 1️⃣5️⃣ Financial Advisor, 1️⃣6️⃣ Loan & Debt Management (Myanmar style: WavePay, KBZ, Pawnshop), 1️⃣7️⃣ Small Business Coach (Tea shop, Online shop, SME), 1️⃣8️⃣ Business Planning, 1️⃣9️⃣ Sales Script Expert (Myanmar style), 2️⃣0️⃣ Online Income Expert (Freelancing, Content, Digital Products), 2️⃣1️⃣ Myanmar Market Price Analyst (Gold, USD, Fuel, Commodities), 2️⃣2️⃣ Broker & Trading, 2️⃣3️⃣ Insurance Agent\n"
    "=== နည်းပညာနှင့် ဒီဂျစ်တယ် (Tech & Digital) ===\n"
    "2️⃣4️⃣ Tech Support, 2️⃣5️⃣ Computer Repair, 2️⃣6️⃣ Phone Repair, 2️⃣7️⃣ Software, 2️⃣8️⃣ Internet & Data Packages (MPT/ATOM), 2️⃣9️⃣ Facebook Page Growth, 3️⃣0️⃣ TikTok Content Coach, 3️⃣1️⃣ YouTube Myanmar Market, 3️⃣2️⃣ Video Editing & Design, 3️⃣3️⃣ Social Media Crisis Management, 3️⃣4️⃣ Content Strategy, 3️⃣5️⃣ Digital Marketing\n"
    "=== ကျန်းမာရေးနှင့် နေ့စဉ်ဘဝ (Health & Daily Life) ===\n"
    "3️⃣6️⃣ Fitness, 3️⃣7️⃣ Nutrition, 3️⃣8️⃣ Health & Medicine (⚠️ No prescriptions), 3️⃣9️⃣ Household Problem Solver (Water pump, Electricity trips), 4️⃣0️⃣ Food & Cooking (Myanmar dishes, cost calc), 4️⃣1️⃣ Mother & Baby care, 4️⃣2️⃣ Beauty & Hair, 4️⃣3️⃣ Pet Care, 4️⃣4️⃣ Gardening & Landscaping, 4️⃣5️⃣ Myanmar Transport Guide (Bus, Train, Express, Delivery), 4️⃣6️⃣ Local Travel Tips\n"
    "=== ဥပဒေ၊ အုပ်ချုပ်မှုနှင့် လူမှုရေး (Legal & Social) ===\n"
    "4️⃣7️⃣ Legal Awareness (Myanmar: Land, Contract, Police report, ⚠️ Disclaimer), 4️⃣8️⃣ Government Documents & Public Services, 4️⃣9️⃣ Work & Migration (Visa, Thailand, Korea, Japan, ⚠️ Scam warning), 5️⃣0️⃣ Job & Career Advisor (CV, Interview, Myanmar job market), 5️⃣1️⃣ HR & Management, 5️⃣2️⃣ Leadership, 5️⃣3️⃣ Business Ethics\n"
    "=== ယဉ်ကျေးမှု၊ အနုပညာနှင့် ဖျော်ဖြေရေး (Culture, Arts & Entertainment) ===\n"
    "5️⃣4️⃣ Myanmar Culture Advisor (Traditions, Manners, Ceremonies), 5️⃣5️⃣ Event Planner (Wedding, Donation, Birthday), 5️⃣6️⃣ Astrology (Myanmar & Western, entertainment only), 5️⃣7️⃣ Traditional Arts (Painting, Blacksmith, Goldsmith, Lacquerware, Marble, Woodcarving – 10 arts), 5️⃣8️⃣ Movies & Cinema (Reviews, legal links), 5️⃣9️⃣ Literature (Poetry, Novels, Writing), 6️⃣0️⃣ Music, 6️⃣1️⃣ Languages & Translation (Myanmar, English, Chinese, Thai), 6️⃣2️⃣ History (Myanmar & World), 6️⃣3️⃣ Religion (Buddhism, Christianity, Islam – neutral), 6️⃣4️⃣ Vipassana Meditation, 6️⃣5️⃣ Traditional Medicine (Herbal, ⚠️ Disclaimer)\n"
    "=== မြန်မာ့ဈေးကွက်နှင့် အထူးပြု (Myanmar-Specific & Niche) ===\n"
    "6️⃣6️⃣ Myanmar Commodity Prices (Rice, Oil, Gold, etc.), 6️⃣7️⃣ Thai Lottery (2D, 3D, etc., ⚠️ Illegal warning), 6️⃣8️⃣ Telegram Bots & Automation, 6️⃣9️⃣ Real Estate & Property (Buy/Sell, Land Titles), 7️⃣0️⃣ Vehicle Trading (Cars, Bikes), 7️⃣1️⃣ Clothing & Fashion, 7️⃣2️⃣ Cosmetics & Skincare, 7️⃣3️⃣ Electronics Trading, 7️⃣4️⃣ Bicycle & Motorcycle Repair, 7️⃣5️⃣ Construction & Electrical, 7️⃣6️⃣ Agriculture & Farming, 7️⃣7️⃣ Livestock & Animal Care, 7️⃣8️⃣ Environment & Wildlife, 7️⃣9️⃣ Natural Disaster Preparedness, 8️⃣0️⃣ Local Foods (Mohannga, Shan noodles, etc.)\n"
    "=== နောက်ထပ် ခေတ်စားနယ်ပယ်များ (Additional Trending Topics) ===\n"
    "8️⃣1️⃣ Personal Branding, 8️⃣2️⃣ Public Speaking, 8️⃣3️⃣ Conflict Resolution, 8️⃣4️⃣ Negotiation Skills, 8️⃣5️⃣ Email & Business Writing, 8️⃣6️⃣ Presentation Skills, 8️⃣7️⃣ Data Analysis, 8️⃣8️⃣ Basic Accounting, 8️⃣9️⃣ Tax Basics (Myanmar), 9️⃣0️⃣ Import/Export Basics, 9️⃣1️⃣ E-commerce (Shopee, Lazada), 9️⃣2️⃣ Dropshipping, 9️⃣3️⃣ Affiliate Marketing, 9️⃣4️⃣ SEO & Blogging, 9️⃣5️⃣ App Development, 9️⃣6️⃣ AI Tools (ChatGPT, Midjourney, etc.), 9️⃣7️⃣ Photography, 9️⃣8️⃣ Cooking Recipes (International), 9️⃣9️⃣ Travel Planning (Domestic & International), 🔟0️⃣ Spiritual Growth (non-religious), 🔟1️⃣ Relationship & Marriage, 🔟2️⃣ Parenting, 🔟3️⃣ Education & Study Tips, 🔟4️⃣ Exam Preparation, 🔟5️⃣ Career Change, 🔟6️⃣ Retirement Planning, 🔟7️⃣ Freelance Platforms (Upwork, Fiverr), 🔟8️⃣ Personal Finance Apps, 🔟9️⃣ Cybersecurity & Online Safety, 🔟🔟 Myanmar Internet Scams (Fraud prevention)\n"
    "=== နေ့စဉ်သုံး နှင့် လူမှုဆက်ဆံရေး (Daily & Social) ===\n"
    "1️⃣1️⃣1️⃣ နေ့စဉ်သုံးစကားများ (Daily Phrases) - အင်္ဂလိပ်၊ မြန်မာ၊ တရုတ် နေ့စဉ်သုံး စကားပြောအသုံးအနှုန်းများ၊ စကားပြောလေ့ကျင့်နည်းများ သင်ကြားပေးခြင်း။\n"
    "1️⃣1️⃣2️⃣ အလာဘသလာဘ (Small Talk) - လူမှုဆက်ဆံရေးတွင် ကျွမ်းကျင်စွာ ပြောဆိုနည်း၊ မိတ်ဆက်စကား၊ စကားစမြည်ပြောနည်းများ သင်ကြားပေးခြင်း။\n\n"

    "=== ပြောင်းလဲနိုင်သော အသံများ (Adaptive Persona Modes) ===\n"
    "အသုံးပြုသူရဲ့ မေးခွန်းနဲ့ လေသံကို လိုက်ပြီး အောက်ပါ ပုံစံတွေထဲက တစ်ခုခုကို အလိုအလျောက် ခံယူနိုင်ပါတယ်။\n"
    "- **Sarcastic Friend Mode** - ရင်းရင်းနှီးနှီး နောက်ပြောင်တဲ့လေသံ (ဥပမာ: 'ဟာ Than... ဒီလိုလုပ်ရင် ငွေမစုဘူးနော် 😏')\n"
    "- **Strict Teacher Mode** - စည်းကမ်းတကျ၊ အမှားကို တိုက်ရိုက်ထောက်ပြ (ဥပမာ: 'မင်း ဒီလိုလုပ်ရင် မအောင်ဘူး၊ discipline လိုတယ်!')\n"
    "- **Grandma Mode** - နွေးထွေးပြီး ဂရုစိုက်တဲ့လေသံ (ဥပမာ: 'ကလေးရေ... ငွေကို သုံးသုံးချွေတာချွေတာနဲ့...')\n"
    "- **Trader Mode** - ဈေးကွက်နဲ့ ပတ်သက်ရင် အေးဆေးတည်ငြိမ်တဲ့ အကြံဉာဏ် (ဥပမာ: 'Bro... market pump လာတယ်, calm down first.')\n"
    "- **Love Guru Mode** - အချစ်ရေး အကြံပေးရင် နူးညံ့သိမ်မွေ့ (ဥပမာ: 'မင်း crush ကို text မပို့ခင် deep breath တစ်ချက်ယူပါ...')\n"
    
    "=== အထူးလမ်းညွှန်ချက်များ (Important Instructions) ===\n"
    "1. ရာသီဥတုမေးရင် - တိကျတဲ့ ဒေသကို မေးပြီးမှ ဖြေပါ။\n"
    "2. ရင်းနှီးမြှုပ်နှံမှု/ခရစ်ပတို/ရွှေ - 'အာမခံအမြတ်' လို့ ဘယ်တော့မှ မပြောပါနဲ့။ အန္တရာယ်ရှိနိုင်ကြောင်း အမြဲသတိပေးပါ။\n"
    "3. ဆေးဝါး - ဆေးညွှန်းမပေးပါနဲ့။ ဆရာဝန်နဲ့ တိုင်ပင်ရန် ပြောပါ။\n"
    "4. ဥပဒေနဲ့ ပတ်သက်ရင် - တရားဝင်အကြံဉာဏ်မဟုတ်ကြောင်း အမြဲထည့်ပြောပါ။\n"
    "5. ဗီဇာ/အလုပ်အကိုင် - Scam အန္တရာယ် သတိပေးပါ၊ တရားဝင်လမ်းကြောင်းကို ညွှန်ပါ။\n"
    "6. ထိုင်းထီ - တရားမဝင်ကြောင်း၊ အန္တရာယ်ရှိကြောင်း ရိုးသားစွာ သတိပေးပါ။\n"
    "7. ဘာသာရေး - ဘက်မလိုက်ဘဲ လေးစားမှုဖြင့် ရှင်းပြပါ။\n"
    "8. ကိုယ်ရေးကိုယ်တာ ပြဿနာများ - စာနာနားလည်မှုနဲ့ ဖြေပါ၊ အကြံပြုချက်များသာ ပေးပါ။\n"
    "9. မြန်မာ့ဈေးနှုန်း - နောက်ဆုံးရအချက်အလက်များကို 'ခန့်မှန်းချက်' အဖြစ်သာ ဖော်ပြပါ။\n"
    "10. ရုပ်ရှင်/စာပေ - တရားဝင်လင့်များသာ အကြံပြုပါ။\n"
    "11. ငွေရေးကြေးရေး - အာမခံချက်မပေးပါနဲ့။ တာဝန်ယူမှုရှိအောင် ပြောပါ။\n"
    "12. နိုင်ငံရေး/သတင်း - ဘက်မလိုက်ဘဲ တည်ငြိမ်စွာ ရှင်းပြပါ။\n"
    "13. သမိုင်းမေးရင် - တိကျသော သမိုင်းအချက်အလက်များကို သာလျှင် ဖြေဆိုပါ။ မသေချာလျှင် 'မသေချာပါ' ဟု ပြောပါ။\n"
    "14. ကုန်စျေးနှုန်းမေးရင် - နောက်ဆုံးရရှိသော အချက်အလက်များကို အသုံးပြုပြီး 'လက်ရှိခန့်မှန်းချက်' ဟု ဖော်ပြပါ။ တိကျသေချာမှုမရှိလျှင် အသိပေးပါ။\n"
    "15. ရုပ်ရှင်မေးရင် - တရားဝင်လင့်များသာ ပေးပါ။ မူပိုင်ခွင့်ကို လေးစားပါ။\n"
    "16. စာပေမေးရင် - စာရေးဆရာများ၏ မူပိုင်ခွင့်ကို လေးစားပါ။ စာပေသဘောတရားများကို ရှင်းလင်းစွာ ဖော်ပြပါ။\n"
    "17. ဘာသာရေးမေးရင် - ဘာသာတစ်ခုချင်းစီ၏ ယုံကြည်မှုကို လေးစားပါ။ ဘက်မလိုက်ပါနှင့်။\n"
    "18. ဝိပဿနာမေးရင် - ဆရာတစ်ဦး၏ လမ်းညွှန်မှုဖြင့် ကျင့်သုံးရန် အကြံပြုပါ။ စိတ်ပိုင်းဆိုင်ရာ ထိခိုက်မှုများကို သတိပြုပါ။\n"
    "19. အဆင့်နိမ့်ပတ်ဝန်းကျင်မေးရင် - စိတ်ဓာတ်ပိုင်း အားပေးမှုနှင့် တကယ့်အကူအညီရှာဖွေရန် ညွှန်ကြားပါ။ လူမှုရေးအကူအညီများကို အကြံပြုပါ။\n"
    "20. ပန်းချီ/ပန်းပဲ/ပန်းထိမ်/ပန်းတမော့ - ရိုးရာအနုပညာများကို ထိန်းသိမ်းရန် အကြံပြုပါ။ အန္တရာယ်ကင်းရှင်းရေးကို အလေးထားပါ။\n"
    "21. ဘတ်စ်ကားလမ်းညွှန်မေးရင် - လက်ရှိအချိန်ဇယားများကို စစ်ဆေးရန် အကြံပြုပါ။ လမ်းကြောင်းပြောင်းလဲမှုများ ရှိနိုင်ကြောင်း သတိပေးပါ။\n"
    "22. ငွေရေးကြေးရေးမေးရင် - ဘတ်ဂျက်ရေးဆွဲနည်း၊ ကြွေးမြီစီမံခန့်ခွဲနည်းများကို ရိုးသားစွာ အကြံပေးပါ။ 'ချမ်းသာအောင် ဒီလိုလုပ်ပါ' မျိုး အာမခံချက် မပေးပါနှင့်။\n"
    "23. ဘာသာပြန်မေးရင် - မူရင်းအဓိပ္ပာယ်ကို တိကျစွာ ထိန်းသိမ်းပြီး လွယ်ကူရှင်းလင်းသော စကားလုံးများဖြင့် ပြန်ဆိုပါ။\n"
    "24. အလုပ်အကိုင်နှင့် ကိုယ်ရေးမှတ်တမ်း (CV) မေးရင် - CV ပြင်ဆင်နည်း၊ Interview အတွက် လေ့ကျင့်နည်း၊ လစာနှုန်းထားများကို လက်တွေ့ကျကျ ပြောပါ။\n"
    "25. ပြည်ပအလုပ်အကိုင်/ဗီဇာမေးရင် - လုပ်ငန်းစဉ်များကို ရှင်းပြပြီး၊ Scam အန္တရာယ်ကို အမြဲ သတိပေးပါ။ တရားဝင်လမ်းကြောင်းသာ ညွှန်ပါ။\n"
    "26. အွန်လိုင်းဝင်ငွေမေးရင် - လက်တွေ့ကျသော ဝင်ငွေခန့်မှန်းချက်နှင့် WavePay/KBZPay ကန့်သတ်ချက်များကို ပြောပါ။\n"
    "27. စီးပွားရေးလုပ်ငန်း/ဆိုင်ဖွင့်မေးရင် - ကုန်ကျစရိတ်၊ ရင်းနှီးမြှုပ်နှံမှုနှင့် Break-even Point (အမြတ်စတင်မည့်နေရာ) ကို တွက်ချက်ပြပါ။\n"
    "28. Facebook/TikTok/YouTube မေးရင် - လူကြိုက်များသော Content ပုံစံ၊ Algorithm နှင့် Monetization ရရှိနိုင်မှု အနေအထားကို ရိုးသားစွာ ပြောပါ။\n"
    "29. အရောင်းအ၀ယ် (Sales) မေးရင် - မြန်မာဖောက်သည်များ၏ စိတ်နေစိတ်ထားကို နားလည်ပြီး Objection ဖြေရှင်းနည်းများ သင်ပေးပါ။\n"
    "30. စိတ်ကျန်းမာရေး/စိတ်ဖိစီးမှုမေးရင် - ယေဘုယျ အားပေးစကားသာ ပြောပါ။ စိတ်ရောဂါကုဆရာဝန်နှင့် တိုင်ပင်ရန် အမြဲ အကြံပြုပါ။\n"
    "31. ဘုန်းကြီးစကား/တရားဓမ္မ ပုံစံမေးရင် - အေးဆေးတည်ငြိမ်ပြီး ဉာဏ်ပညာနှင့် ပြောပါ။ ဘာသာရေးအသွင် မဆောင်ပါစေနှင့်။\n"
    "32. ဆိုရှယ်မီဒီယာ အကျပ်အတည်း (Page Hack, Fake Account) မေးရင် - Report လုပ်နည်း၊ Appeal တင်နည်းအဆင့်ဆင့်ကို ရှင်းပြပါ။\n"
    "33. အကြွေးနှင့် ချေးငွေ စီမံခန့်ခွဲရာတွင် - WavePay/KBZ ချေးငွေများ၊ ပေါင်ဆိုင်အန္တရာယ်နှင့် Debt Snowball Method (အကြွေးချေနည်းစနစ်) ကို ရှင်းပြပါ။\n"
    "34. မြန်မာ့ဈေးနှုန်း (ဆီ၊ ရွှေ၊ ဒေါ်လာ) မေးရင် - 'ခန့်မှန်းချက်' သာဖြစ်ကြောင်း အမြဲဖော်ပြပါ။\n"
    "35. မြန်မာ့ယဉ်ကျေးမှု (အလှူ၊ မင်္ဂလာ၊ ဓလေ့) မေးရင် - လေးစားမှုဖြင့် ရှင်းပြပါ။\n"
    "36. ပွဲစီစဉ်ခြင်း (Wedding, Donation, Birthday) မေးရင် - အစီအစဉ်၊ ဘတ်ဂျက်နှင့် ပြင်ဆင်ရမည့်အချက်များကို ကူညီစီစဉ်ပေးပါ။\n"
    "37. အိမ်မှုကိစ္စ (ရေပန့်ပျက်၊ မီးလိုင်းပြတ်) မေးရင် - လွယ်ကူသော ပြုပြင်နည်းများနှင့် အန္တရာယ်ကင်းရှင်းရေးကို ဦးစားပေး ပြောပါ။\n"
    "38. ဥပဒေအခြေခံ (မြေယာ၊ စာချုပ်၊ ရဲစခန်း) မေးရင် - Disclaimer ဖြင့်သာ ဖြေပါ။ ရှေ့နေနှင့် တိုင်ပင်ရန် အမြဲ ညွှန်ပါ။\n"
    "39. ခရီးသွားလမ်းညွှန် (ဘတ်စ်၊ ရထား၊ Express) မေးရင် - လက်ရှိ အချိန်ဇယားများကို စစ်ဆေးရန် ပြောပါ။\n"
    "40. မြန်မာအစားအစာ (မုန့်ဟင်းခါး၊ ရှမ်းခေါက်ဆွဲ) ချက်ပြုတ်ရောင်းချလိုပါက - ကုန်ကျစရိတ်တွက်ချက်နည်းကို ထည့်သွင်း ပြောပါ။\n"
    "41. ဖုန်း/အင်တာနက် Package မေးရင် - MPT/ATOM/Telenor အကောင်းဆုံး Plan များကို နှိုင်းယှဉ်ပြပါ။\n"
    "42. မြန်မာဗေဒင် (နေ့၊ လ၊ နှစ်) မေးရင် - ဖျော်ဖြေမှုသဘောသာ (Entertainment only) ဖြစ်ကြောင်း ပြောပြီး ဖြေပါ။\n"
    "43. Bonus Modes (Sarcastic, Strict Teacher, Grandma, Trader, Love Guru) - အသုံးပြုသူ၏ မေးခွန်းပုံစံပေါ် မူတည်ပြီး သင့်လျော်သော လေသံ (Tone) ကို အလိုအလျောက် ခံယူပါ။\n"
    "44. စီးပွားရေးနှင့် ပတ်သက်သော ခန့်မှန်းချက်များ - လက်တွေ့ကျပြီး တိကျသော အချက်အလက်များသာ ပေးပါ။ တကယ့်လုပ်ငန်းရှင်များနှင့် တိုင်ပင်သင့်ကြောင်း အမြဲ သတိပေးပါ။\n"
    "45. နေ့စဉ်သုံးစကားများ သင်ကြားရာတွင် - အသုံးပြုသူ၏ အဆင့်ကို စစ်ဆေးပြီး ရိုးရှင်းသော စကားစုများမှ စတင်ပါ။\n"
    "46. အလာဘသလာဘ ပြောဆိုရာတွင် - လူမှုဆက်ဆံရေး ယဉ်ကျေးမှုကို လေးစားပြီး၊ ဖော်ရွေမှုကို ဦးစားပေးပါ။\n\n"

    "အသုံးပြုသူရဲ့ မေးခွန်းကို ကျွမ်းကျင်သူတစ်ယောက်လို ဖြေပေးရမယ်။\n"
    "အဖြေတွေကို မြန်မာလိုပဲ ပြန်ရမယ်။ လေးလေးနက်နက်၊ ရယ်စရာ၊ မိတ်ဆွေလို ပြောရမယ်။\n"
    "ဥပဒေ၊ ကျန်းမာရေး၊ ငွေကြေးဆိုင်ရာ အကြံပြုချက်များသည် အထွေထွေ အချက်အလက်သာဖြစ်ပြီး ကျွမ်းကျင်သူများနှင့် တိုင်ပင်ရန် သတိပေးရမယ်။"
)


# ====== AI Model ======
async def ask_model(prompt: str, user_id: str = None) -> str:
    user_context = ""
    cache_allowed = True
    cache_key = f"{prompt.strip()}|{user_id}"

    if user_id:
        user_data = get_user(user_id)
        if user_data:
            (plan, usage, proof_status, _, _, _, goals, weaknesses, dream, career, money_mindset, relationship, birthdate) = user_data
            if any([goals, weaknesses, dream, career, money_mindset, relationship, birthdate]):
                cache_allowed = False
                user_context = "\n\n[အသုံးပြုသူ၏ ကိုယ်ရေးအချက်အလက်များ]\n"
                if goals: user_context += f"- ပန်းတိုင်: {goals}\n"
                if career: user_context += f"- အလုပ်အကိုင်: {career}\n"
                if dream: user_context += f"- အိပ်မက်: {dream}\n"
                if weaknesses: user_context += f"- အားနည်းချက်: {weaknesses}\n"
                if money_mindset: user_context += f"- ငွေကြေးစိတ်ဓာတ်: {money_mindset}\n"
                if relationship: user_context += f"- ဆက်ဆံရေး: {relationship}\n"
                if birthdate: user_context += f"- မွေးနေ့: {birthdate}\n"

    if cache_allowed:
        cached = get_cached_response(cache_key)
        if cached:
            return cached

    # DeepSeek Retry
    for attempt in range(2):
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.post(
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY.strip()}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "deepseek-chat",
                        "messages": [
                            {"role": "system", "content": system_prompt + user_context},
                            {"role": "user", "content": prompt}
                        ],
                        "max_tokens": 800,
                        "temperature": 0.85
                    },
                )
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    final_answer = result["choices"][0]["message"]["content"].strip()
                    if cache_allowed:
                        save_cached_response(cache_key, final_answer)
                    return final_answer
                elif "error" in result:
                    raise Exception(result["error"]["message"])
                else:
                    raise Exception("Unexpected API response: " + str(result))
        except Exception as e:
            if attempt == 1:
                logger.error(f"DeepSeek failed: {e}. Falling back to GPT-4o-mini.")
                break
            await asyncio.sleep(2 ** attempt)

    # GPT-4o-mini Retry
    for attempt in range(2):
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.post(
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY.strip()}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "gpt-4o-mini",
                        "messages": [
                            {"role": "system", "content": system_prompt + user_context},
                            {"role": "user", "content": prompt}
                        ],
                        "max_tokens": 800,
                        "temperature": 0.85
                    },
                )
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    final_answer = result["choices"][0]["message"]["content"].strip()
                    if cache_allowed:
                        save_cached_response(cache_key, final_answer)
                    return final_answer
                elif "error" in result:
                    raise Exception(result["error"]["message"])
                else:
                    raise Exception("Unexpected API response: " + str(result))
        except Exception as e:
            if attempt == 1:
                raise e
            await asyncio.sleep(2 ** attempt)

async def send_daily_coaching(bot):
    try:
        conn = sqlite3.connect("bot_users.db", check_same_thread=False)
        c = conn.cursor()
        c.execute("SELECT user_id FROM users WHERE plan IN ('free', 'premium_plus')")
        users = c.fetchall()
        conn.close()

        if not users:
            return

        prompt = "မင်္ဂလာပါ၊ ဒီနေ့အတွက် နေ့စဉ် ဘဝလမ်းညွှန်စကား (Daily Coaching Message) ကို မစ္စတာသန်း (Mr.T) ရဲ့ အသံနဲ့ ရေးပါ။ ယုံကြည်မှု၊ အလုပ်အကိုင်၊ ငွေကြေးအတွေးအခေါ် အကြောင်းတွေ ပါစေ။"
        message = await ask_model(prompt)
        if not message or len(message) < 10:
            message = "🌅 ဒီနေ့အတွက် အကောင်းဆုံး နေ့တစ်နေ့ ဖြစ်ပါစေ။"
        
        sent = 0
        for user in users:
            try:
                await bot.send_message(chat_id=int(user[0]), text=f"🔥 **Daily Coaching Message**\n\n{message}")
                sent += 1
                await asyncio.sleep(0.05)
            except Exception as e:
                logger.error(f"Failed to send daily coaching to {user[0]}: {e}")
        logger.info(f"✅ Daily coaching sent to {sent} users.")
    except Exception as e:
        logger.error(f"❌ Daily coaching error: {e}")

# ====== Scheduler ======
def run_scheduler(bot):
    schedule.every(30).days.do(reset_usage)
    
    mm_tz = pytz.timezone("Asia/Yangon")
    now_mm = datetime.now(mm_tz)
    
    schedule.every().day.at("08:00").do(lambda: asyncio.run(send_daily_coaching(bot)))
    schedule.every().day.at("03:00").do(lambda: backup_and_send(bot))
    
    logger.info("⏰ Scheduler started (Myanmar Time).")
    while True:
        schedule.run_pending()
        time.sleep(60)

# ====== Rate Limiting for verifyid ======
verify_attempts = {}

# ====== Auto Verify Payment ======
async def auto_verify_payment(user_id: str, tx_id: str, plan: str, bot) -> bool:
    try:
        conn = get_db_connection()
        c = conn.cursor()

        c.execute("SELECT used, expiry_date FROM transactions WHERE tx_id=? AND user_id=?", (tx_id, user_id))
        row = c.fetchone()
        if not row:
            conn.close()
            return False
        
        used, expiry_date = row
        if used == 1:
            conn.close()
            return False
        
        if expiry_date and datetime.utcnow().isoformat() > expiry_date:
            conn.close()
            return False

        c.execute("UPDATE transactions SET used=1 WHERE tx_id=? AND user_id=?", (tx_id, user_id))
        price = PLAN_LIMITS[plan]["price"]
        c.execute("UPDATE users SET plan=?, proof_status='approved', usage_count=0, price=? WHERE user_id=?",
                  (plan, price, user_id))
        conn.commit()
        conn.close()
        
        # Admin notification
        for admin_id in ADMIN_IDS:
            try:
                await bot.send_message(
                    chat_id=admin_id,
                    text=f"📋 Auto-verify completed!\nUser: {user_id}\nPlan: {plan}\nCode: {tx_id}"
                )
            except Exception as e:
                logger.error(f"Admin notification error: {e}")
        
        logger.info(f"✅ Auto-verified: user={user_id}, tx={tx_id}, plan={plan}")
        return True
    except Exception as e:
        logger.error(f"❌ Auto verify error: {e}")
        return False

# ====== Admin: Generate Transaction Code (Auto-Send to User) ======
async def gen_code(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        return await update.message.reply_text("❌ Admin only.")
    
    if len(context.args) < 2:
        return await update.message.reply_text("Usage: /gen_code <user_id> <plan>")
    
    target_user = context.args[0]
    plan = context.args[1].lower()
    
    if plan not in PLAN_LIMITS:
        return await update.message.reply_text("❌ Invalid plan.")
    
    # Generate unique 5-digit code
    import random
    tx_id = None
    for _ in range(10):
        candidate = f"{random.randint(0, 99999):05d}"
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("SELECT tx_id FROM transactions WHERE tx_id=?", (candidate,))
        exists = c.fetchone()
        conn.close()
        if not exists:
            tx_id = candidate
            break
    
    if not tx_id:
        return await update.message.reply_text("❌ Failed to generate unique code. Try again.")
    
    # 1 day expiry
    expiry_date = (datetime.utcnow() + timedelta(days=1)).isoformat()
    
    # Save to database
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("INSERT INTO transactions (tx_id, user_id, plan, timestamp, used, expiry_date) VALUES (?, ?, ?, ?, ?, ?)",
              (tx_id, target_user, plan, datetime.utcnow().isoformat(), 0, expiry_date))
    conn.commit()
    conn.close()
    
    # ✅ Inline Keyboard with Verify Button
    keyboard = [[InlineKeyboardButton("✅ Verify Code", callback_data=f"verify_{tx_id}")]]
    reply_markup = InlineKeyboardMarkup(keyboard)
    
    # Send code to user with verify button
    try:
        await context.bot.send_message(
            chat_id=int(target_user),
            text=f"🎫 **သင့် {plan.upper()} Plan အတွက် Code ပါ။**\n\n"
                 f"🔑 Code: `{tx_id}`\n"
                 f"⏰ သက်တမ်း: ၂၄ နာရီ\n\n"
                 f"အောက်ပါ ခလုတ်ကို နှိပ်ပြီး Plan ရယူပါ။",
            reply_markup=reply_markup
        )
        await update.message.reply_text(f"✅ Code `{tx_id}` ကို User `{target_user}` ဆီ ပို့ပြီးပါပြီ။")
    except Exception as e:
        await update.message.reply_text(
            f"⚠️ User `{target_user}` ဆီ ပို့လို့မရပါဘူး။\n"
            f"Code: `{tx_id}` ကို ကိုယ်တိုင်ပို့ပေးပါ။"
        )

# ====== User Self-Verify ======
async def verifyid(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)

    if len(context.args) < 1:
        return await update.message.reply_text("Usage: /verifyid <transaction_code>")

    tx_id = context.args[0]

    if not re.match(r"^\d{5}$", tx_id):
        return await update.message.reply_text("❌ Invalid transaction code. Must be exactly 5 digits.")

    # Rate limiting
    now = time.time()
    if user_id not in verify_attempts:
        verify_attempts[user_id] = []
    verify_attempts[user_id] = [t for t in verify_attempts[user_id] if now - t < 600]
    
    if len(verify_attempts[user_id]) >= 5:
        return await update.message.reply_text(
            "❌ Too many attempts. Please wait 10 minutes and try again."
        )
    verify_attempts[user_id].append(now)

    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT plan, expiry_date FROM transactions WHERE tx_id=? AND user_id=?", (tx_id, user_id))
    row = c.fetchone()

    if not row:
        conn.close()
        return await update.message.reply_text("❌ Transaction code not found or not assigned to you.")

    plan, expiry_date = row
    
    if expiry_date and datetime.utcnow().isoformat() > expiry_date:
        conn.close()
        return await update.message.reply_text("❌ This transaction code has expired. Please contact admin.")

    success = await auto_verify_payment(user_id, tx_id, plan, context.bot)
    
    if success:
        await update.message.reply_text(
            f"✅ Transaction {tx_id} verified.\n"
            f"🎉 Your plan upgraded to {plan}."
        )
    else:
        await update.message.reply_text("❌ Transaction code already used or error occurred.")

# ====== Telegram Bot Command Handlers ======
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if context.args:
        ref_code = context.args[0]
        if ref_code.startswith("REF"):
            inviter_id = ref_code.replace("REF", "")
            if inviter_id != user_id:
                give_referral_reward(inviter_id, user_id)

    user = get_user(user_id)
    if not user:
        add_user(user_id, "free")

    keyboard = [
        [InlineKeyboardButton("📌 Plan ရွေးရန်", callback_data="start_plan")],
        [InlineKeyboardButton("ℹ️ လမ်းညွှန်ချက်များ (Help)", callback_data="start_help")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await update.message.reply_text(
        "🙏 မင်္ဂလာပါ။ ကျွန်တော် မစ္စတာသန်းပါ။\n"
        "သင့်ရဲ့ လက်ထောက် အဖြစ်နဲ့ ကိုယ်ရေးကိုယ်တာ၊ အလုပ်အကိုင်နဲ့ "
        "တခြားလုပ်ဆောင်ရမယ့် အရာတွေကို ယုံကြည်စွာ ဖြေရှင်းပေးဖို့ အသင့်ပါဗျ။\n\n"
        "အောက်ပါခလုတ်များမှ ရွေးချယ်ပါ။",
        reply_markup=reply_markup
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📌 အသုံးပြုနည်း:\n"
        "/start - စတင်ရန်\n"
        "/help - အကူအညီ\n"
        "/subscribe - Plan ရွေးရန်\n"
        "/ask <q> - မေးရန်\n"
        "/status - အနေအထား\n"
        "/profile - ကိုယ်ရေးမှတ်တမ်း\n"
        "/habit - အလေ့အထ\n"
        "/referral - ဖိတ်ရန်\n"
        "/gen_code - Admin code generator\n"
        "/verifyid - User self-verify\n\n"
        "🎯 ကျွန်တော် အကြံပေးနိုင်တဲ့ နယ်ပယ် ၁၂၀+ ခုရှိပါတယ်။"
    )

async def referral(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    ref_code = generate_ref_code(user_id)
    bot_username = (await context.bot.get_me()).username
    ref_link = f"https://t.me/{bot_username}?start={ref_code}"
    inviter_count = get_referral_count(user_id)

    await update.message.reply_text(
        f"🔗 **သင့် Referral Link**\n\n"
        f"`{ref_link}`\n\n"
        f"📊 သင်ဖိတ်ကြားထားပြီးသား သူငယ်ချင်း အရေအတွက်: **{inviter_count}**\n\n"
        f"🎁 **ရမည့်အခွင့်အရေး:**\n"
        f"👉 သူငယ်ချင်းတစ်ယောက် ဒီလင့်ကနေ ဝင်သုံးတိုင်း **သင် ၅၀ ကြိမ် အပိုသုံးခွင့် ရပါမယ်။**\n"
        f"👉 ဒါမှမဟုတ် ပိုမိုသုံးစွဲလိုပါက အောက်ပါ Plan များကို ဝယ်ယူနိုင်ပါတယ်:\n"
        f"⭐ **Basic (7,000 MMK)**\n"
        f"💎 **Premium (20,000 MMK)**\n"
        f"👑 **Premium+ (35,000 MMK)**\n\n"
        f"📌 မျှဝေပြီး အပိုသုံးခွင့်ရယူလိုက်ပါ!"
    )

async def subscribe(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    plan = context.args[0] if context.args else "free"
    if plan not in PLAN_LIMITS:
        allowed = ", ".join(PLAN_LIMITS.keys())
        await update.message.reply_text(f"❌ '{plan}' မရှိပါ။ ရနိုင်တဲ့ Plan: {allowed}")
        return

    keyboard = [
        [InlineKeyboardButton("📌 Free (အခမဲ့)", callback_data="sub_free")],
        [InlineKeyboardButton("⭐ Basic (7,000 MMK)", callback_data="sub_basic")],
        [InlineKeyboardButton("💎 Premium (20,000 MMK)", callback_data="sub_premium")],
        [InlineKeyboardButton("👑 Premium+ (35,000 MMK)", callback_data="sub_premium_plus")],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await update.message.reply_text(
        "❌ **သင့် Free Plan ၏ သုံးခွင့် ကုန်သွားပါပြီ။**\n\n"
        "🚀 ဆက်လက်သုံးစွဲရန် အောက်ပါ Plan များထဲမှ တစ်ခုကို ရွေးချယ်ပါ:\n"
        "⭐ **Basic (7,000 MMK)**\n"
        "💎 **Premium (20,000 MMK)**\n"
        "👑 **Premium+ (35,000 MMK)**\n\n"
        "📸 ငွေလွှဲပြီး Proof ဓာတ်ပုံ ပို့ပေးပါ။",
        reply_markup=reply_markup
    )

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    user_id = str(query.from_user.id)
    data = query.data

    # ... (ရှိပြီးသား start_help, sub_, verify_, send_code_)

    if data == "start_plan":
        keyboard = [
           [InlineKeyboardButton("📌 Free (အခမဲ့)", callback_data="sub_free_1m")],
           [InlineKeyboardButton("⭐ Basic (7,000 MMK/လ)", callback_data="show_price_basic")],
           [InlineKeyboardButton("💎 Premium (20,000 MMK/လ)", callback_data="show_price_premium")],
           [InlineKeyboardButton("👑 Premium+ (35,000 MMK/လ)", callback_data="show_price_premium_plus")],
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await query.edit_message_text(
            "📌 Plan ရွေးချယ်ပါ။\n\n"
            "💡 သက်တမ်းကြာလေ ဈေးလျှော့လေ ရမှာပါ။",
            reply_markup=reply_markup
        )
        return

    # ====== Show Price (Plan) ======
    if data.startswith("show_price_"):
        plan = data.replace("show_price_", "")
        if plan not in PLAN_LIMITS:
            await query.edit_message_text("❌ Invalid plan.")
            return
        
        prices = PLAN_LIMITS[plan]["prices"]
        keyboard = [
            [InlineKeyboardButton(f"📆 ၁ လ ({prices['1m']:,} MMK)", callback_data=f"sub_{plan}_1m")],
            [InlineKeyboardButton(f"📆 ၃ လ ({prices['3m']:,} MMK) - လျှော့ဈေး", callback_data=f"sub_{plan}_3m")],
            [InlineKeyboardButton(f"📆 ၁၂ လ ({prices['12m']:,} MMK) - အများဆုံးလျှော့ဈေး", callback_data=f"sub_{plan}_12m")],
            [InlineKeyboardButton("🔙 နောက်သို့", callback_data="start_plan")]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await query.edit_message_text(
            f"📌 **{plan.upper()} Plan** အတွက် သက်တမ်းရွေးပါ။\n\n"
            f"💰 ၁ လ: {prices['1m']:,} MMK\n"
            f"💰 ၃ လ: {prices['3m']:,} MMK (လျှော့ဈေး)\n"
            f"💰 ၁၂ လ: {prices['12m']:,} MMK (အများဆုံးလျှော့ဈေး)",
            reply_markup=reply_markup
        )
        return

    # ====== Subscribe with Duration ======
    if data.startswith("sub_"):
        parts = data.split("_", 2)
        if len(parts) < 3:
            await query.edit_message_text("❌ Invalid request.")
            return
        plan = parts[1]
        duration = parts[2]
        
        if plan not in PLAN_LIMITS:
            await query.edit_message_text("❌ Invalid plan.")
            return
        if duration not in ["1m", "3m", "12m"]:
            await query.edit_message_text("❌ Invalid duration.")
            return
        
        price = get_plan_price(plan, duration)
        expiry_date = calculate_expiry(duration)
        
        conn = sqlite3.connect("bot_users.db", check_same_thread=False)
        c = conn.cursor()
        c.execute("UPDATE users SET plan=?, proof_status='waiting', price=?, duration=?, expiry_date=? WHERE user_id=?",
                  (plan, price, duration, expiry_date, user_id))
        conn.commit()
        conn.close()
        
        duration_label = get_duration_label(duration)
        
        await query.edit_message_text(
            f"📌 **{plan.upper()}** Plan ကို **{duration_label}** အတွက် ရွေးလိုက်ပါပြီ။\n"
            f"💰 စျေးနှုန်း: {price:,} MMK\n"
            f"⏰ သက်တမ်းကုန်ဆုံးရက်: {expiry_date[:10]}\n\n"
            f"📸 ကျေးဇူးပြုပြီး ငွေသွင်း proof screenshot ကို ပို့ပါ။\n\n"
            f"{PAYMENT_INFO}"
        )
        return

    if data == "start_help":
        await query.edit_message_text(
            "📌 အသုံးပြုနည်း:\n"
            "/start - စတင်ရန်\n"
            "/help - အကူအညီ\n"
            "/subscribe - Plan ရွေးရန်\n"
            "/ask <q> - မေးရန်\n"
            "/status - အနေအထား\n"
            "/profile - ကိုယ်ရေးမှတ်တမ်း\n"
            "/habit - အလေ့အထ\n"
            "/referral - ဖိတ်ရန်\n\n"
            "🎯 ကျွန်တော် အကြံပေးနိုင်တဲ့ နယ်ပယ် ၁၂၀+ ခုရှိပါတယ်။"
        )
        return

    if data.startswith("sub_"):
        plan = data.replace("sub_", "")
        if plan not in PLAN_LIMITS:
            await query.edit_message_text("❌ Invalid plan.")
            return

        conn = sqlite3.connect("bot_users.db", check_same_thread=False)
        c = conn.cursor()
        c.execute("UPDATE users SET plan=?, proof_status='waiting', price=? WHERE user_id=?", (plan, PLAN_LIMITS[plan]["price"], user_id))
        conn.commit()
        conn.close()

        price_mmk = PLAN_LIMITS[plan]["price"]
        price_usd = get_price_usd(price_mmk)

        text = (
            f"📌 **{plan}** Plan ကို ရွေးလိုက်ပါပြီ။\n"
            f"💰 စျေးနှုန်း: {price_mmk:,} MMK (~${price_usd}) / month\n\n"
            f"📸 ကျေးဇူးပြုပြီး ငွေသွင်း proof screenshot ကို ပို့ပါ။\n\n"
            f"{PAYMENT_INFO}"
        )
        await query.edit_message_text(text)

    if data.startswith("verify_"):
        tx_id = data.replace("verify_", "")
        
        # Check if code exists and is valid
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("SELECT plan, used, expiry_date FROM transactions WHERE tx_id=? AND user_id=?", (tx_id, user_id))
        row = c.fetchone()
        
        if not row:
            await query.edit_message_text("❌ Code မရှိပါဘူး။")
            conn.close()
            return
        
        plan, used, expiry_date = row
        
        if used == 1:
            await query.edit_message_text("❌ ဒီ Code က သုံးပြီးသားပါ။")
            conn.close()
            return
        
        if expiry_date and datetime.utcnow().isoformat() > expiry_date:
            await query.edit_message_text("❌ ဒီ Code ရဲ့ သက်တမ်းကုန်သွားပါပြီ။")
            conn.close()
            return
        
        # Mark as used and upgrade plan
        c.execute("UPDATE transactions SET used=1 WHERE tx_id=? AND user_id=?", (tx_id, user_id))
        price = PLAN_LIMITS[plan]["price"]
        c.execute("UPDATE users SET plan=?, proof_status='approved', usage_count=0, price=? WHERE user_id=?",
                  (plan, price, user_id))
        conn.commit()
        conn.close()
        
        # Send confirmation to user
        await query.edit_message_text(
            f"🎉 **Plan အဆင့်မြှင့်ပြီးပါပြီ။**\n\n"
            f"📌 Plan: **{plan}**"
        )
        
        # Notify admin
        for admin_id in ADMIN_IDS:
            try:
                await context.bot.send_message(
                    chat_id=admin_id,
                    text=f"📋 User {user_id} က Code `{tx_id}` နဲ့ Plan ကိုယ်တိုင်အသက်သွင်းလိုက်ပါပြီ။\n"
                         f"📌 Plan: {plan}"
                )
            except Exception as e:
                logger.error(f"Admin notification error: {e}")
        return

    if data.startswith("send_code_"):
        if user_id not in [str(a) for a in ADMIN_IDS]:
            await query.edit_message_text("❌ Admin only.")
            return
        
        # data က "send_code_123456789_premium" လိုမျိုး ဖြစ်မယ်
        parts = data.split("_", 2)
        if len(parts) < 3:
            await query.edit_message_text("❌ Invalid request.")
            return
        
        target_user = parts[1]
        plan = parts[2]
        
        # Generate unique 5-digit code
        import random
        tx_id = None
        for _ in range(10):
            candidate = f"{random.randint(0, 99999):05d}"
            conn = get_db_connection()
            c = conn.cursor()
            c.execute("SELECT tx_id FROM transactions WHERE tx_id=?", (candidate,))
            exists = c.fetchone()
            conn.close()
            if not exists:
                tx_id = candidate
                break
        
        if not tx_id:
            await query.edit_message_text("❌ Code ထုတ်လို့မရပါဘူး။ နောက်တစ်ခါ ကြိုးစားပါ။")
            return
        
        # 1 day expiry
        expiry_date = (datetime.utcnow() + timedelta(days=1)).isoformat()
        
        # Save to database
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("INSERT INTO transactions (tx_id, user_id, plan, timestamp, used, expiry_date) VALUES (?, ?, ?, ?, ?, ?)",
                  (tx_id, target_user, plan, datetime.utcnow().isoformat(), 0, expiry_date))
        conn.commit()
        conn.close()
        
        # Send code to user
        try:
            await context.bot.send_message(
                chat_id=int(target_user),
                text=f"🎫 **သင့် {plan.upper()} Plan အတွက် Code ပါ။**\n\n"
                     f"🔑 Code: `{tx_id}`\n"
                     f"⏰ သက်တမ်း: ၂၄ နာရီ\n\n"
                     f"📌 /verifyid {tx_id} လို့ရိုက်ပြီး Plan ရယူပါ။"
            )
            await query.edit_message_text(
                f"✅ Code `{tx_id}` ကို User `{target_user}` ဆီ ပို့ပြီးပါပြီ။\n"
                f"📌 Plan: {plan}"
            )
        except Exception as e:
            await query.edit_message_text(
                f"⚠️ User `{target_user}` ဆီ ပို့လို့မရပါဘူး။\n"
                f"Code: `{tx_id}` ကို ကိုယ်တိုင်ပို့ပေးပါ။"
            )
        
        # Update user's proof_status
        conn = get_db_connection()
        c = conn.cursor()
        c.execute("UPDATE users SET proof_status='pending' WHERE user_id=?", (target_user,))
        conn.commit()
        conn.close()
        return

async def status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    user = get_user(user_id)
    if not user:
        add_user(user_id, "free")
        user = ("free", 0, "none", None, 0, None, None, None, None, None, None, None, None)
    (plan, usage, proof_status, _, price, _, goals, weaknesses, dream, career, money_mindset, relationship, birthdate) = user
    limit = PLAN_LIMITS[plan]["limit"]
    remaining = limit - usage
    price_usd = get_price_usd(price)
    ref_count = get_referral_count(user_id)
    level = usage // 100 + 1
    level_title = "🥉 Bronze" if level == 1 else "🥈 Silver" if level == 2 else "🥇 Gold" if level >= 3 else "🌱 Beginner"
    
    profile_preview = ""
    if goals: profile_preview += f"\n• 🎯 Goals: {goals}"
    if career: profile_preview += f"\n• 💼 Career: {career}"
    if money_mindset: profile_preview += f"\n• 💰 Money: {money_mindset}"
    if dream: profile_preview += f"\n• 🌟 Dream: {dream}"
    if weaknesses: profile_preview += f"\n• ⚠️ Weaknesses: {weaknesses}"
    if relationship: profile_preview += f"\n• ❤️ Relationship: {relationship}"
    if birthdate: profile_preview += f"\n• 🎂 Birthdate: {birthdate}"
    await update.message.reply_text(f"📊 **Your Status**\n🏅 Level: {level_title} (Lv.{level})\n📌 Plan: **{plan}**\n📈 Usage: {usage}/{limit}\n🔋 Remaining: **{remaining}**\n🔍 Proof Status: {proof_status}\n👥 Referrals: **{ref_count}**{profile_preview}")

async def ask(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    chat_type = update.effective_chat.type

    if chat_type in ["group", "supergroup"]:
        await update.message.reply_text(
            "⚠️ /ask command ကို Group ထဲမှာ တိုက်ရိုက်မသုံးပါနဲ့ဗျာ။\n"
            "🔒 ကိုယ်ရေးကိုယ်တာ မေးခွန်းတွေအတွက် Bot ရဲ့ Private Chat (DM) မှာ အသုံးပြုပေးပါ။\n"
            "ဒါမှမဟုတ် Group ထဲမှာ ကျွန်တော့်နာမည် ဒါမှမဟုတ် @BotUsername ကို ခေါ်ပြီး မေးနိုင်ပါတယ်။"
        )
        return

    if not check_limit(user_id):
        keyboard = [[InlineKeyboardButton("⭐ Plan ရွေးရန်", callback_data="start_plan")]]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message.reply_text(
            "❌ **သင့် Free Plan ၏ သုံးခွင့် ကုန်သွားပါပြီ။**\n\n"
            "🚀 ဆက်လက်သုံးစွဲရန် အောက်ပါ Plan များထဲမှ တစ်ခုကို ရွေးချယ်ပါ:\n"
            "⭐ **Basic (7,000 MMK)**\n"
            "💎 **Premium (20,000 MMK)**\n"
            "👑 **Premium+ (35,000 MMK)**\n\n"
            "📸 ငွေလွှဲပြီး Proof ဓာတ်ပုံ ပို့ပေးပါ။",
            reply_markup=reply_markup
        )
        return

    if not context.args:
        await update.message.reply_text("❌ မေးခွန်းထည့်ပေးပါ။\nUsage: /ask <your question>")
        return

    question = " ".join(context.args)
    local_answer = get_local_response(question)
    if local_answer:
        increment_usage(user_id)
        await update.message.reply_text(local_answer)
        return

    await update.message.reply_text("🤔 စဉ်းစားနေပါတယ်...")
    try:
        answer = await ask_model(question, user_id)
        answer = clean_text(answer)
    except Exception as e:
        logger.error(f"Ask command error: {e}")
        await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")
        return
    increment_usage(user_id)
    await update.message.reply_text(answer, disable_web_page_preview=True)

async def profile(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if not context.args:
        await update.message.reply_text("Usage: /profile <field> : <value>\nExample: /profile goals : ကိုယ်ပိုင်လုပ်ငန်းဖွင့်မယ်\nAllowed fields: goals, weaknesses, dream, career, money_mindset, relationship, birthdate")
        return
    text = " ".join(context.args)
    if ":" not in text:
        await update.message.reply_text("❌ Format: field : value")
        return
    field_raw, value = [p.strip() for p in text.split(":", 1)]
    allowed_fields = ["goals", "weaknesses", "dream", "career", "money_mindset", "relationship", "birthdate"]
    if field_raw not in allowed_fields:
        await update.message.reply_text("❌ field မမှန်ပါ။ ခွင့်ပြုထားတဲ့ field တွေသာ ထည့်ပါ။")
        return
    field_map = {"goals": "goals", "weaknesses": "weaknesses", "dream": "dream", "career": "career", "money_mindset": "money_mindset", "relationship": "relationship", "birthdate": "birthdate"}
    key = field_map.get(field_raw)
    update_profile(user_id, key, value)
    await update.message.reply_text(f"✅ `{key}` ကို update လုပ်ပြီးပါပြီ။")

async def habit(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if not context.args:
        await update.message.reply_text("Usage: /habit <habit>\nExample: /habit နေ့တိုင်း ၁၅ မိနစ် စာဖတ်မယ်")
        return
    habit_text = " ".join(context.args)
    add_habit(user_id, habit_text)
    await update.message.reply_text(f"✅ Habit အသစ် ထည့်ပြီးပါပြီ:\n\n• {habit_text}")

async def myhabits(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    habits = get_habits(user_id)
    if not habits:
        await update.message.reply_text("😅 သင့် habits စာရင်း မရှိသေးပါ။")
        return
    text = "📝 **Your Habits (Latest 20)**\n"
    for habit_text, created_at in habits:
        text += f"\n• {habit_text}  ({created_at})"
    await update.message.reply_text(text)

async def proof(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("📸 ငွေသွင်း proof screenshot ကို ဒီ chat ထဲမှာ Photo အနေနဲ့ ပို့ပေးပါ။")

async def sum_numbers(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Usage: /sum <စာရင်း သို့မဟုတ် ကိန်းဂဏန်းများ>")
        return
    text = " ".join(context.args)
    numbers = re.findall(r'\d+(?:\.\d+)?', text)
    if not numbers:
        await update.message.reply_text("❌ ဂဏန်းတွေ မတွေ့ပါဘူး။")
        return
    total = sum(float(num) for num in numbers)
    if total.is_integer():
        total = int(total)
    await update.message.reply_text(f"🧮 ပေါင်းလဒ်စုစုပေါင်း = {total}")

async def read_photo_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['photo_mode'] = 'read'
    await update.message.reply_text("📸 ပုံကို ပို့ပေးပါ။ ဒီပုံထဲက စာသားတွေကို ဖတ်ပေးပါမယ်။")

async def photo_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if not update.message.photo:
        return
    
    # ====== Check if it's OCR mode ======
    mode = context.user_data.get('photo_mode', 'proof')
    
    if mode == 'read':
        context.user_data['photo_mode'] = 'proof'
        file = await context.bot.get_file(update.message.photo[-1].file_id)
        img_bytes = await file.download_as_bytearray()
        b64_image = base64.b64encode(img_bytes).decode('utf-8')
        
        await update.message.reply_text("🤔 ပုံထဲက စာသားတွေကို ဖတ်နေပါတယ်...")
        
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.post(
                    OPENROUTER_URL,
                    headers={"Authorization": f"Bearer {OPENROUTER_API_KEY.strip()}", "Content-Type": "application/json"},
                    json={
                        "model": "gpt-4o-mini",
                        "messages": [
                            {"role": "system", "content": "You are an OCR assistant. Extract all text from the image."},
                            {"role": "user", "content": [
                                {"type": "text", "text": "Read the text in this image."},
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_image}"}}
                            ]}
                        ],
                        "max_tokens": 800
                    }
                )
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    text = result["choices"][0]["message"]["content"].strip()
                    await update.message.reply_text(text)
                else:
                    await update.message.reply_text("❌ ဖတ်လို့မရပါဘူး။")
        except Exception as e:
            await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")
        return
    
    # ====== Proof Mode (Auto-Verify System) ======
    file_id = update.message.photo[-1].file_id
    
    # Save to database (still track proof for record)
    update_proof(user_id, file_id)
    
    # User confirmation
    await update.message.reply_text(
        "📸 Screenshot ကို လက်ခံရရှိပါပြီ။\n\n"
        "💡 သင့် Plan ကို အသက်သွင်းရန် ကျေးဇူးပြုပြီး အောက်ပါအတိုင်း လုပ်ဆောင်ပါ:\n"
        "1️⃣ Admin ထံမှ Code ကို ရယူပါ\n"
        "2️⃣ /verifyid <code> လို့ရိုက်ထည့်ပါ"
    )
    
    # ====== ✅ Admin Notification (For info only) ======
    for admin_id in ADMIN_IDS:
        try:
            await context.bot.send_message(
                chat_id=admin_id,
                text=f"📋 User `{user_id}` က Screenshot တင်လိုက်ပါပြီ။\n"
                     f"(Auto-Verify System သုံးနေတာမို့ အတည်ပြုစရာမလိုပါဘူး)"
            )
            await context.bot.send_photo(
                chat_id=admin_id,
                photo=file_id,
                caption=f"📸 Proof from User {user_id}\n\n"
                        f"✅ ဒီ User အတွက် Code ထုတ်ပေးရန်:\n"
                        f"/gen_code {user_id} <plan>"
            )
        except Exception as e:
            logger.error(f"Admin notification error for {admin_id}: {e}")

async def image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Usage: /image <ပုံဖော်ပြချက်>")
        return
    
    prompt = " ".join(context.args)
    await update.message.reply_text("🎨 ပုံဖန်တီးနေပါတယ်... ခဏစောင့်ပါ...")
    
    try:
        url = f"https://image.pollinations.ai/prompt/{prompt}"
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.get(url)
            if response.status_code == 200:
                await update.message.reply_photo(photo=response.content)
            else:
                await update.message.reply_text(f"❌ ပုံထုတ်လို့မရပါဘူး: {response.status_code}")
    except Exception as e:
        await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")

async def make_cv(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text(
            "Usage: /cv <နာမည်> | <ဖုန်း> | <ပညာရေး> | <အတွေ့အကြုံ>\n"
            "Example: /cv ကိုအောင် | 09xxxx | တက္ကသိုလ် | Sales အတွေ့အကြုံ ၂ နှစ်"
        )
        return

    data = " ".join(context.args).split("|")
    if len(data) < 4:
        await update.message.reply_text("❌ ပုံစံမှားနေပါတယ်။ `|` နဲ့ ခြားပြီး အချက် ၄ ချက် ထည့်ပါ။")
        return

    name = data[0].strip()
    phone = data[1].strip()
    education = data[2].strip()
    experience = data[3].strip()

    doc = Document()
    doc.add_heading(f"CV - {name}", 0)
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(experience)

    filename = "CV_Output.docx"
    doc.save(filename)

    with open(filename, "rb") as f:
        await update.message.reply_document(document=f, filename=filename, caption=f"📄 {name} ၏ CV ဖိုင် ပါ။")
    os.remove(filename)

async def get_weather(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Usage: /weather <မြို့နယ်နာမည်>\nExample: /weather Yangon")
        return

    city = " ".join(context.args)
    await update.message.reply_text(f"🌍 {city} ရဲ့ ရာသီဥတုကို ရှာနေပါတယ်...")

    try:
        geo_res = await httpx.AsyncClient(timeout=10).get(
            f"https://geocoding-api.open-meteo.com/v1/search?name={city}&count=1&language=en&format=json"
        )
        geo_data = geo_res.json()
        if not geo_data.get("results"):
            await update.message.reply_text("❌ ဒီမြို့ကို ရှာမတွေ့ပါဘူး။ နာမည်အပြည့်အစုံ ရိုက်ပါ။")
            return

        lat = geo_data["results"][0]["latitude"]
        lon = geo_data["results"][0]["longitude"]
        place_name = geo_data["results"][0]["name"]

        weather_res = await httpx.AsyncClient(timeout=10).get(
            f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current_weather=true"
        )
        weather_data = weather_res.json()
        current = weather_data.get("current_weather", {})

        temperature = current.get("temperature")
        wind_speed = current.get("windspeed")
        weather_code = current.get("weathercode")

        conditions = {
            0: "သာယာနေပါတယ် ☀️", 1: "နည်းနည်းတိမ်ထူနေပါတယ် ⛅", 2: "တိမ်ထူနေပါတယ် ☁️", 3: "အုံ့ဆိုင်းနေပါတယ် 🌥️",
            45: "မြူဆိုင်းနေပါတယ် 🌫️", 51: "အလွန်နည်းသော မိုးစက်ကျနေပါတယ် 🌦️", 61: "မိုးရွာနေပါတယ် 🌧️",
            63: "မိုးအသင့်အတင့် ရွာနေပါတယ် 🌧️", 65: "မိုးသည်းထန်စွာ ရွာနေပါတယ် ⛈️", 71: "နှင်းကျနေပါတယ် ❄️",
            80: "မိုးသက်မုန်တိုင်း ကျနေပါတယ် 🌩️", 95: "မိုးကြိုးမုန်တိုင်း ကျနေပါတယ် ⚡"
        }
        condition_text = conditions.get(weather_code, "ရာသီဥတု အခြေအနေ သိမရပါဘူး")

        await update.message.reply_text(
            f"📍 **{place_name}** ရဲ့ လက်ရှိရာသီဥတု - \n\n"
            f"🌡️ အပူချိန်: {temperature}°C\n"
            f"💨 လေတိုက်နှုန်း: {wind_speed} km/h\n"
            f"☁️ အခြေအနေ: {condition_text}"
        )
    except Exception as e:
        await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")

async def text_to_speech(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Usage: /tts <စာသား> (အသံထုတ်ချင်တဲ့ စာသားကို ရိုက်ပါ)")
        return

    text = " ".join(context.args)
    await update.message.reply_text("🔊 အသံဖိုင်ကို ဖန်တီးနေပါတယ်...")

    try:
        tts = gTTS(text=text, lang='my', slow=False)
        filename = "speech.mp3"
        tts.save(filename)

        with open(filename, "rb") as f:
            await update.message.reply_audio(audio=f, title="Mr.T ဖတ်ပေးတဲ့ အသံ", caption="🔊 အသံဖိုင် ပါ။")
        os.remove(filename)
    except Exception as e:
        await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")

async def fetch_news(update: Update, context: ContextTypes.DEFAULT_TYPE, topic: str = None):
    await update.message.reply_text("📰 သတင်းတွေ ရှာနေပါတယ်... ခဏစောင့်ပါ...")
    try:
        if topic:
            url = f"https://news.google.com/rss/search?q={topic}&hl=en&gl=MM&ceid=MM:en"
        else:
            url = "https://news.google.com/rss?hl=en&gl=MM&ceid=MM:en"
        
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url)
            response.raise_for_status()
            root = ET.fromstring(response.text)
        
        items = root.findall(".//item")[:10]
        if not items:
            await update.message.reply_text("❌ သတင်းမတွေ့ပါဘူး။ နောက်တစ်ခါ ထပ်ကြိုးစားပါ။")
            return
        
        news_text = "📰 **နောက်ဆုံးရ သတင်းများ**\n\n"
        for i, item in enumerate(items, 1):
            title = item.find("title").text
            link = item.find("link").text
            news_text += f"{i}. {title}\n"
            news_text += f"   🔗 {link}\n\n"
        
        await update.message.reply_text(news_text, disable_web_page_preview=True)
    except Exception as e:
        await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")

async def news_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.args:
        topic = " ".join(context.args)
        await fetch_news(update, context, topic)
    else:
        await fetch_news(update, context)

# ====== Admin Commands ======
async def pending_requests(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        return await update.message.reply_text("❌ Admin only.")
    
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT user_id, plan FROM users WHERE proof_status='waiting' AND plan != 'free'")
    results = c.fetchall()
    conn.close()
    
    if not results:
        return await update.message.reply_text("📭 လက်ရှိ Pending Requests မရှိပါဘူး။")
    
    keyboard = []
    for uid, plan in results:
        button_text = f"📤 Send Code to {uid} ({plan})"
        callback_data = f"send_code_{uid}_{plan}"
        keyboard.append([InlineKeyboardButton(button_text, callback_data=callback_data)])
    
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "📋 **Pending Plan Requests**\n\n"
        "အောက်ပါ User တွေအတွက် Code ထုတ်ပေးရန် ခလုတ်ကို နှိပ်ပါ။",
        reply_markup=reply_markup
    )
    
async def verify(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        await update.message.reply_text("❌ Admin only.")
        return
    if len(context.args) < 2:
        await update.message.reply_text("Usage: /verify <user_id> <plan>")
        return
    target_id, plan = context.args[0], context.args[1]
    if plan not in PLAN_LIMITS:
        await update.message.reply_text("❌ Invalid plan.")
        return
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("UPDATE users SET plan=?, proof_status='approved', price=? WHERE user_id=?", (plan, PLAN_LIMITS[plan]["price"], target_id))
    conn.commit()
    conn.close()
    await update.message.reply_text(f"✅ User {target_id} ကို {plan} plan သို့ verify လုပ်ပြီးပါပြီ။")

async def pending_proofs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        await update.message.reply_text("❌ Admin only.")
        return
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT user_id, plan, proof_file_id, proof_timestamp FROM users WHERE proof_status='pending'")
    rows = c.fetchall()
    conn.close()
    if not rows:
        await update.message.reply_text("✅ Pending proofs မရှိပါ။")
        return
    text = "📋 **Pending Proofs**\n"
    for uid, plan, file_id, ts in rows:
        text += f"\n• User: {uid} | Plan: {plan} | File: {file_id} | Time: {ts}"
    await update.message.reply_text(text)

async def approve_proof(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        await update.message.reply_text("❌ Admin only.")
        return
    if not context.args:
        await update.message.reply_text("Usage: /approve_proof <user_id>")
        return
    target_id = context.args[0]
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT plan FROM users WHERE user_id=?", (target_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        await update.message.reply_text("❌ User not found.")
        return
    plan = row[0]
    price = PLAN_LIMITS[plan]["price"]
    c.execute("UPDATE users SET proof_status='approved', usage_count=0, price=? WHERE user_id=?", (price, target_id))
    conn.commit()
    conn.close()
    await update.message.reply_text(f"✅ Proof for user {target_id} approved. Plan: {plan}")

async def reject_proof(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        await update.message.reply_text("❌ Admin only.")
        return
    if not context.args:
        await update.message.reply_text("Usage: /reject_proof <user_id>")
        return
    target_id = context.args[0]
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("UPDATE users SET proof_status='rejected' WHERE user_id=?", (target_id,))
    conn.commit()
    conn.close()
    await update.message.reply_text(f"❌ Proof for user {target_id} rejected.")

async def broadcast(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = str(update.effective_user.id)
    if user_id not in [str(a) for a in ADMIN_IDS]:
        await update.message.reply_text("❌ Admin only.")
        return
    if not context.args:
        await update.message.reply_text("Usage: /broadcast <message>")
        return
    message = " ".join(context.args)
    conn = sqlite3.connect("bot_users.db", check_same_thread=False)
    c = conn.cursor()
    c.execute("SELECT user_id FROM users")
    users = c.fetchall()
    conn.close()
    sent = 0
    for (uid,) in users:
        try:
            await context.bot.send_message(chat_id=int(uid), text=message)
            sent += 1
            await asyncio.sleep(0.03)
        except Exception as e:
            logger.error(f"Broadcast failed to {uid}: {e}")
    await update.message.reply_text(f"✅ Broadcast sent to {sent} users.")

# ====== Main Message Handler ======
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.text:
        return
    user_id = str(update.effective_user.id)
    chat_type = update.effective_chat.type
    text = update.message.text

    if text.lower() in ["free", "basic", "premium", "premium_plus"]:
        context.args = [text.lower()]
        await subscribe(update, context)
        return

    # ====== Group Chat ======
    if chat_type in ["group", "supergroup"]:
        if await is_bot_mentioned(update, context):
            # Remove bot name
            for name in BOT_NAMES:
                if name.lower() in text.lower():
                    text = re.sub(name, "", text, flags=re.IGNORECASE)
            bot_username = (await context.bot.get_me()).username
            if bot_username:
                text = re.sub(f"@{bot_username}", "", text, flags=re.IGNORECASE)
            text = text.strip()
            if not text:
                return
            
            if not check_limit(user_id):
                await update.message.reply_text("❌ သုံးခွင့်ကုန်သွားပါပြီ။ Plan အသစ်ရွေးပါ။")
                return
            
            # ✅ Local Response (Token မကုန်စေရန်)
            local_answer = get_local_response(text)
            if local_answer:
                await update.message.reply_text(local_answer)
                return
            
            # ✅ AI Response (Token ကုန်မယ်)
            await update.message.reply_text("🤔 စဉ်းစားနေပါတယ်...")
            try:
                short_prompt = f"Group chat ဖြစ်လို့ တိုတိုနဲ့ ဖြေပါ။ {text}"
                answer = await ask_model(short_prompt, user_id)
                answer = clean_text(answer)
            except Exception as e:
                logger.error(f"Group message error: {e}")
                await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")
                return
            increment_usage(user_id)
            await update.message.reply_text(answer, disable_web_page_preview=True)
            return

    # ====== Private Chat ======
    if chat_type == "private":
        if not check_limit(user_id):
            keyboard = [[InlineKeyboardButton("⭐ Plan ရွေးရန်", callback_data="start_plan")]]
            reply_markup = InlineKeyboardMarkup(keyboard)
            await update.message.reply_text(
                "❌ **သင့် Free Plan ၏ သုံးခွင့် ကုန်သွားပါပြီ။**\n\n"
                "🚀 ဆက်လက်သုံးစွဲရန် အောက်ပါ Plan များထဲမှ တစ်ခုကို ရွေးချယ်ပါ:\n"
                "⭐ **Basic (7,000 MMK)**\n"
                "💎 **Premium (20,000 MMK)**\n"
                "👑 **Premium+ (35,000 MMK)**\n\n"
                "📸 ငွေလွှဲပြီး Proof ဓာတ်ပုံ ပို့ပေးပါ။",
                reply_markup=reply_markup
            )
            return

        if "သတင်း" in text or "news" in text.lower():
            await fetch_news(update, context)
            return
        
        # ✅ Local Response (Token မကုန်စေရန်)
        local_answer = get_local_response(text)
        if local_answer:
            await update.message.reply_text(local_answer)
            return
        
        # ✅ AI Response (Token ကုန်မယ်)
        await update.message.reply_text("🤔 စဉ်းစားနေပါတယ်...")
        try:
            answer = await ask_model(text, user_id)
            answer = clean_text(answer)
        except Exception as e:
            logger.error(f"Handle message error: {e}")
            await update.message.reply_text(f"⚠️ Error: {str(e)[:100]}")
            return
        increment_usage(user_id)
        bot_name = get_bot_name(text)
        if bot_name:
            answer = f"{bot_name} ပြောတယ်... {answer}"
        await update.message.reply_text(answer, disable_web_page_preview=True)

async def is_bot_mentioned(update: Update, context: ContextTypes.DEFAULT_TYPE) -> bool:
    if not update.message or not update.message.text:
        return False
    text = update.message.text.lower()
    for name in BOT_NAMES:
        if name.lower() in text:
            return True
    bot_username = (await context.bot.get_me()).username
    if bot_username and f"@{bot_username.lower()}" in text:
        return True
    return False

# ====== Main ======
def main():
    init_db()
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    # User Commands
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("subscribe", subscribe))
    application.add_handler(CommandHandler("status", status))
    application.add_handler(CommandHandler("ask", ask))
    application.add_handler(CommandHandler("profile", profile))
    application.add_handler(CommandHandler("habit", habit))
    application.add_handler(CommandHandler("myhabits", myhabits))
    application.add_handler(CommandHandler("proof", proof))
    application.add_handler(CommandHandler("referral", referral))
    application.add_handler(CommandHandler("sum", sum_numbers))
    application.add_handler(CommandHandler("cv", make_cv))
    application.add_handler(CommandHandler("weather", get_weather))
    application.add_handler(CommandHandler("tts", text_to_speech))
    application.add_handler(CommandHandler("readphoto", read_photo_command))
    application.add_handler(CommandHandler("image", image))
    application.add_handler(CommandHandler("news", news_command))
    
    # ✅ New Auto-Verify Commands
    application.add_handler(CommandHandler("gen_code", gen_code))      # Admin code generator
    application.add_handler(CommandHandler("verifyid", verifyid))      # User self-verify
    
    # Admin Commands
    application.add_handler(CommandHandler("pending_requests", pending_requests))  
    application.add_handler(CommandHandler("verify", verify))
    application.add_handler(CommandHandler("pending_proofs", pending_proofs))
    application.add_handler(CommandHandler("approve_proof", approve_proof))
    application.add_handler(CommandHandler("reject_proof", reject_proof))
    application.add_handler(CommandHandler("broadcast", broadcast))
    application.add_handler(CommandHandler("gen_code", gen_code))
    application.add_handler(CommandHandler("verifyid", verifyid))
    
    # Callback & Message Handlers
    application.add_handler(CallbackQueryHandler(button_handler))
    application.add_handler(MessageHandler(filters.PHOTO & (~filters.COMMAND), photo_handler))
    application.add_handler(MessageHandler(filters.TEXT & (~filters.COMMAND), handle_message))

    def run_flask():
        flask_app.run(host="0.0.0.0", port=5000)

    threading.Thread(target=run_flask, daemon=True).start()
    threading.Thread(target=run_scheduler, args=(application.bot,), daemon=True).start()

    application.run_polling()

if __name__ == "__main__":
    main()
